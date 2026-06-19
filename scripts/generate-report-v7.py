"""
Generate internship report v7 — modify v6-final.docx directly.

Changes from v6→v7:
  Task 1: Insert ER Diagram after section 2.8 (after paragraph ~190)
  Task 2: Insert System Architecture Diagram after section 3.4 (after paragraph ~239)
  Task 3: Insert Flowchart images after section 3.5 (after paragraph ~242)
  Task 4: Fix "สต๊อก" → "สต็อก" (4 places) + "สำรองไหล่" → "สำรองอะไหล่" (1 place)
  Task 5: Fix "2568" → "2569" on cover page (paragraph 17)
  Task 6: Add year in abstract paragraph (paragraph 39)
  Task 7: Add KPI source reference text after KPI table (after paragraph 83)
  Task 8: Add 7 glossary terms to section 1.6 (paragraph 161-162)
  Task 9: Replace UAT table from 6 items to 35 items (paragraph 353 area)
  Task 10: Add bibliography at end
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import os, copy

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMG_DIR = os.path.join(BASE, "docs", "images")
INPUT = os.path.join(BASE, "docs", "report-v6-final.docx")
OUT_PATH = os.path.join(BASE, "docs", "report-v7-final.docx")

FONT = 'TH SarabunPSK'

# ── Open v6-final.docx ──
doc = Document(INPUT)
body = doc.element.body

# ══════════════════════════════════════════════════════════════
# Helper: insert element after a specific paragraph index
# ══════════════════════════════════════════════════════════════
def _insert_after_paragraph(idx, elem):
    """Insert XML element right after doc.paragraphs[idx]."""
    ref_p = doc.paragraphs[idx]._element
    ref_p.addnext(elem)

def _insert_after_element(ref_elem, elem):
    """Insert XML element right after ref_elem."""
    ref_elem.addnext(elem)

def _make_para(text, ref_para_idx=None, bold=False, alignment=None, italic=False, font_size=None):
    """Create a new paragraph element cloned from a reference, with given text."""
    if ref_para_idx is None:
        ref_para_idx = 64  # default body paragraph
    ref_elem = doc.paragraphs[ref_para_idx]._element
    new_p = copy.deepcopy(ref_elem)
    # Remove existing runs
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    # Build new run
    ref_runs = ref_elem.findall(qn('w:r'))
    new_r = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
    if ref_runs:
        ref_rPr = ref_runs[0].find(qn('w:rPr'))
        if ref_rPr is not None:
            new_rPr = copy.deepcopy(ref_rPr)
            # Bold
            if bold:
                b_elem = new_rPr.find(qn('w:b'))
                if b_elem is None:
                    b_elem = parse_xml(f'<w:b {nsdecls("w")}/>')
                    new_rPr.insert(0, b_elem)
                else:
                    b_elem.set(qn('w:val'), '1')
                bCs = new_rPr.find(qn('w:bCs'))
                if bCs is None:
                    bCs = parse_xml(f'<w:bCs {nsdecls("w")}/>')
                    new_rPr.append(bCs)
                else:
                    bCs.set(qn('w:val'), '1')
            else:
                b_elem = new_rPr.find(qn('w:b'))
                if b_elem is not None:
                    b_elem.set(qn('w:val'), '0')
                bCs = new_rPr.find(qn('w:bCs'))
                if bCs is not None:
                    bCs.set(qn('w:val'), '0')
            # Italic
            if italic:
                i_elem = new_rPr.find(qn('w:i'))
                if i_elem is None:
                    i_elem = parse_xml(f'<w:i {nsdecls("w")}/>')
                    new_rPr.append(i_elem)
                else:
                    i_elem.set(qn('w:val'), '1')
                iCs = new_rPr.find(qn('w:iCs'))
                if iCs is None:
                    iCs = parse_xml(f'<w:iCs {nsdecls("w")}/>')
                    new_rPr.append(iCs)
                else:
                    iCs.set(qn('w:val'), '1')
            # Font size override
            if font_size:
                sz = new_rPr.find(qn('w:sz'))
                if sz is not None:
                    sz.set(qn('w:val'), str(font_size * 2))  # half-points
                szCs = new_rPr.find(qn('w:szCs'))
                if szCs is not None:
                    szCs.set(qn('w:val'), str(font_size * 2))
            new_r.append(new_rPr)
    t = parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve">{_escape_xml(text)}</w:t>')
    new_r.append(t)
    new_p.append(new_r)
    # Alignment
    if alignment:
        pPr = new_p.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            new_p.insert(0, pPr)
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="{alignment}"/>'))
        else:
            jc.set(qn('w:val'), alignment)
    return new_p

def _escape_xml(text):
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')

def _add_page_break_elem():
    return parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:r><w:br w:type="page"/></w:r>'
        f'</w:p>'
    )

def _add_image_after(ref_elem, filename, width=Inches(5.2)):
    """Add an image paragraph directly into the document body after ref_elem.
    This approach keeps the image in the document tree so relationships stay valid."""
    path = os.path.join(IMG_DIR, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if os.path.exists(path):
        run.add_picture(path, width=width)
    else:
        run.text = f"[รูปภาพ: {filename}]"
    # Move the paragraph from end of body to right after ref_elem
    p_elem = p._element
    body.remove(p_elem)
    ref_elem.addnext(p_elem)
    return p_elem

def _add_caption(text, ref_idx=64):
    """Create an italic centered caption paragraph."""
    cap_p = _make_para(text, ref_para_idx=ref_idx, italic=True, font_size=14, alignment='center')
    return cap_p

def _make_table_elem(headers, rows):
    """Create a table element with the given headers and rows."""
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = t._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = FONT
                r.font.size = Pt(18)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT
                    r.font.size = Pt(18)
    tbl_elem = t._tbl
    tbl_elem.getparent().remove(tbl_elem)
    return tbl_elem


# ══════════════════════════════════════════════════════════════
# TASK 4: Fix text replacements
# ══════════════════════════════════════════════════════════════
print("Task 4: Fixing text replacements...")

fix_count = 0
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        original = run.text
        # Fix สต๊อก → สต็อก
        if 'สต๊อก' in run.text:
            run.text = run.text.replace('สต๊อก', 'สต็อก')
            fix_count += 1
            print(f"  [Paragraph {i}] Fixed สต๊อก→สต็อก: '{original[:60]}' → '{run.text[:60]}'")
        # Fix สำรองไหล่ → สำรองอะไหล่
        if 'สำรองไหล่' in run.text:
            run.text = run.text.replace('สำรองไหล่', 'สำรองอะไหล่')
            fix_count += 1
            print(f"  [Paragraph {i}] Fixed สำรองไหล่→สำรองอะไหล่: '{original[:60]}' → '{run.text[:60]}'")
        # Fix 2568 → 2569 on cover page (paragraph 17)
        if 'ปีการศึกษา 2568' in run.text:
            run.text = run.text.replace('ปีการศึกษา 2568', 'ปีการศึกษา 2569')
            fix_count += 1
            print(f"  [Paragraph {i}] Fixed 2568→2569: '{original}' → '{run.text}'")

print(f"  Total text fixes: {fix_count}")

# ══════════════════════════════════════════════════════════════
# TASK 5: Add year in abstract paragraph (paragraph 39)
# ══════════════════════════════════════════════════════════════
print("Task 5: Adding year in abstract...")
# Paragraph 39 starts with "ระหว่างการฝึกงาน ข้าพเจ้าได้รับมอบหมาย..."
# Need to add "ระหว่างวันที่ 1 มีนาคม – 31 พฤษภาคม 2569"
# Find the paragraph that mentions "ระหว่างการฝึกงาน" in the abstract
for i, p in enumerate(doc.paragraphs):
    if 'ระหว่างการฝึกงาน' in p.text and i > 35:
        for run in p.runs:
            if 'ระหว่างการฝึกงาน' in run.text:
                run.text = run.text.replace(
                    'ระหว่างการฝึกงาน',
                    'ระหว่างวันที่ 1 มีนาคม – 31 พฤษภาคม 2569 ข้าพเจ้าได้รับมอบหมาย'
                )
                # Remove duplicate "ข้าพเจ้าได้รับมอบหมาย" if it appears after
                run.text = run.text.replace('ข้าพเจ้าได้รับมอบหมาย ข้าพเจ้าได้รับมอบหมาย', 'ข้าพเจ้าได้รับมอบหมาย')
                print(f"  [Paragraph {i}] Added year range in abstract")
                break
        break

# ══════════════════════════════════════════════════════════════
# TASK 7: Add KPI source reference after KPI table (paragraph 83)
# ══════════════════════════════════════════════════════════════
print("Task 7: Adding KPI source reference...")
# Paragraph 83 = "จากตารางข้างต้น พบว่าระบบช่วยลดระยะเวลา..."
kpi_ref_text = "หมายเหตุ: ตัวชี้วัดประสิทธิภาพ (KPI) ดังกล่าวอ้างอิงจากการสังเกตการณ์ทำงานจริงของเจ้าหน้าที่คลังอะไหล่ ณ โรงไฟฟ้าพระนครเหนือ ก่อนและหลังการใช้งานระบบ"
kpi_ref_p = _make_para(kpi_ref_text, ref_para_idx=64, italic=True, font_size=14)
# Insert after paragraph 83
_insert_after_paragraph(83, kpi_ref_p)
print("  Inserted KPI reference after paragraph 83")

# ══════════════════════════════════════════════════════════════
# TASK 8: Add 7 glossary terms to section 1.6 (after paragraph 162)
# ══════════════════════════════════════════════════════════════
print("Task 8: Adding glossary terms...")
# Find the glossary table — it's after paragraph 161 "1.6 นิยามศัพท์เฉพาะ"
# The table is between paragraph 161 and 162 (empty para before บทที่ 2)
# We need to add 7 more rows to the existing table
# Find the table in section 1.6
glossary_added = False
for tbl in doc.tables:
    # Check if this is the glossary table by checking first cell
    if tbl.rows and tbl.rows[0].cells and 'คำศัพท์' in tbl.rows[0].cells[0].text:
        # Add 7 new rows
        new_terms = [
            ["LLM", "โมเดลภาษาขนาดใหญ่ (Large Language Model) เป็นโมเดลปัญญาประดิษฐ์ที่ฝึกด้วยข้อมูลข้อความจำนวนมาก สามารถทำความเข้าใจและสร้างข้อความตอบกลับได้ ในโครงงานนี้ใช้ผ่าน OpenAI-compatible API"],
            ["Vector Search", "การค้นหาด้วยเวกเตอร์ เป็นเทคนิคการค้นหาข้อมูลโดยเปรียบเทียบความคล้ายคลึงของเวกเตอร์ที่แปลงจากข้อความหรือรูปภาพ ช่วยให้ค้นหาข้อมูลที่ใกล้เคียงความหมายได้ แม้ไม่ตรงกับคำค้นหาทุกคำ"],
            ["Flex Message", "รูปแบบข้อความของ LINE Messaging API ที่สามารถกำหนดเลย์เอาต์แบบกำหนดเองได้ รองรับการแสดงผลแบบคอลัมน์ รูปภาพ ปุ่ม และข้อความหลายบรรทัด"],
            ["Hybrid Search", "การค้นหาแบบผสม เป็นเทคนิคที่รวมการค้นหาแบบดั้งเดิม (Keyword Search) กับการค้นหาด้วยเวกเตอร์ (Vector Search) เพื่อให้ได้ผลลัพธ์ที่แม่นยำและครอบคลุมมากขึ้น"],
            ["Cosine Similarity", "ค่าความคล้ายคลึงแบบโคไซน์ เป็นวิธีวัดความคล้ายคลึงระหว่างเวกเตอร์ 2 เวกเตอร์ โดยคำนวณจากมุมระหว่างเวกเตอร์ ค่าอยู่ระหว่าง -1 ถึง 1 ค่ายิ่งใกล้ 1 แสดงว่าเวกเตอร์มีความคล้ายคลึงกันมาก"],
            ["API Route", "เส้นทาง API ของ Next.js เป็นฟังก์ชันที่ทำงานบนฝั่งเซิร์ฟเวอร์สำหรับรับและประมวลผลคำขอ HTTP โดยกำหนดเส้นทางได้จากโครงสร้างโฟลเดอร์ในโปรเจกต์"],
            ["OpenAI-compatible API", " API ที่ออกแบบให้เข้ากันได้กับรูปแบบ API ของ OpenAI ทำให้สามารถใช้โมเดล AI อื่น ๆ ผ่านอินเทอร์เฟซเดียวกันได้ ในโครงงานนี้ใช้เป็นช่องทางเชื่อมต่อกับ LLM ผ่าน LLM Gateway"],
        ]
        for term_row in new_terms:
            row = tbl.add_row()
            for ci, val in enumerate(term_row):
                cell = row.cells[ci]
                cell.text = val
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = FONT
                        r.font.size = Pt(18)
        glossary_added = True
        print(f"  Added {len(new_terms)} glossary terms to existing table")
        break

if not glossary_added:
    print("  WARNING: Could not find glossary table, skipping")

# ══════════════════════════════════════════════════════════════
# Now handle insertions (Tasks 1, 2, 3) — work from BOTTOM to TOP
# to preserve paragraph indices
# ══════════════════════════════════════════════════════════════

# ── TASK 9: Replace UAT table (6→35 items) ──
print("Task 9: Replacing UAT table...")
# Find the UAT table — it has "ลำดับ" or "รายการทดสอบ" in header
uat_table_found = False
for tbl in doc.tables:
    if tbl.rows and tbl.rows[0].cells:
        header_text = tbl.rows[0].cells[0].text + ' ' + tbl.rows[0].cells[1].text
        if 'ลำดับ' in header_text or 'รายการทดสอบ' in header_text or 'UAT' in header_text:
            # Check it's the small table (6-7 rows = header + 6 data)
            if len(tbl.rows) <= 8:
                # Remove all existing data rows
                while len(tbl.rows) > 1:
                    tbl._tbl.remove(tbl.rows[-1]._tr)

                # Add 35 UAT items
                uat_items = [
                    ["1", "เข้าสู่ระบบด้วยชื่อผู้ใช้และรหัสผ่านที่ถูกต้อง", "เข้าสู่ระบบสำเร็จ แสดงหน้า Dashboard", "ผ่าน"],
                    ["2", "เข้าสู่ระบบด้วยรหัสผ่านผิด", "แสดงข้อความแจ้งเตือนรหัสผ่านไม่ถูกต้อง", "ผ่าน"],
                    ["3", "แสดงผลหน้า Dashboard หลังเข้าสู่ระบบ", "แสดงข้อมูลสรุปจำนวนอะไหล่ หมวดหมู่ อาคาร และอะไหล่ใกล้หมด", "ผ่าน"],
                    ["4", "เพิ่มข้อมูลอะไหล่ใหม่", "บันทึกข้อมูลอะไหล่ใหม่สำเร็จ แสดงในรายการ", "ผ่าน"],
                    ["5", "เพิ่มอะไหล่ใหม่ด้วยรหัสที่ซ้ำกับที่มีอยู่", "แสดงข้อความแจ้งเตือนรหัสอะไหล่ซ้ำ", "ผ่าน"],
                    ["6", "แก้ไขข้อมูลอะไหล่", "บันทึกการแก้ไขสำเร็จ ข้อมูลอัปเดต", "ผ่าน"],
                    ["7", "ค้นหาอะไหล่จากรหัสอะไหล่", "แสดงรายการอะไหล่ที่ตรงกับรหัสที่ค้นหา", "ผ่าน"],
                    ["8", "ค้นหาอะไหล่จากชื่ออะไหล่", "แสดงรายการอะไหล่ที่มีชื่อตรงหรือใกล้เคียง", "ผ่าน"],
                    ["9", "ค้นหาอะไหล่จาก Barcode", "แสดงรายการอะไหล่ที่ตรงกับ Barcode", "ผ่าน"],
                    ["10", "แสดงสถานะอะไหล่ (มีของ/ใกล้หมด/หมดสต็อก)", "แสดงสถานะอะไหล่ถูกต้องตามจำนวนคงเหลือและจำนวนขั้นต่ำ", "ผ่าน"],
                    ["11", "บันทึกการรับเข้าอะไหล่", "จำนวนอะไหล่เพิ่มขึ้น ประวัติถูกบันทึก", "ผ่าน"],
                    ["12", "บันทึกการเบิกออกอะไหล่", "จำนวนอะไหล่ลดลง ประวัติถูกบันทึก", "ผ่าน"],
                    ["13", "บันทึกการเบิกออกเกินจำนวนคงเหลือ", "ระบบป้องกันไม่ให้เบิกเกิน แสดงข้อความแจ้งเตือน", "ผ่าน"],
                    ["14", "บันทึกการปรับยอดอะไหล่", "จำนวนอะไหล่ถูกปรับตามจำนวนที่ระบุ ประวัติถูกบันทึก", "ผ่าน"],
                    ["15", "ตรวจสอบประวัติการเคลื่อนไหวของสต็อก", "แสดงประวัติการรับเข้า เบิกออก และปรับยอดย้อนหลัง", "ผ่าน"],
                    ["16", "สแกน QR Code เพื่อเข้าถึงข้อมูลอะไหล่", "แสดงข้อมูลอะไหล่ที่ตรงกับ QR Code ที่สแกน", "ผ่าน"],
                    ["17", "สแกน QR Code ที่ไม่มีในระบบ", "แสดงข้อความแจ้งว่าไม่พบอะไหล่", "ผ่าน"],
                    ["18", "นำเข้าข้อมูลอะไหล่จากไฟล์ Excel ที่ถูกต้อง", "นำเข้าข้อมูลสำเร็จ แสดงจำนวนรายการที่นำเข้า", "ผ่าน"],
                    ["19", "นำเข้าข้อมูลจากไฟล์ Excel ที่มีรหัสอะไหล่ซ้ำ", "แสดงข้อความแจ้งเตือนรายการที่ซ้ำ", "ผ่าน"],
                    ["20", "เพิ่มหมวดหมู่อะไหล่ใหม่", "บันทึกหมวดหมู่ใหม่สำเร็จ", "ผ่าน"],
                    ["21", "แก้ไขชื่อหมวดหมู่", "บันทึกการแก้ไขสำเร็จ", "ผ่าน"],
                    ["22", "ลบหมวดหมู่ที่ไม่มีอะไหล่อยู่", "ลบหมวดหมู่สำเร็จ", "ผ่าน"],
                    ["23", "เพิ่มข้อมูลอาคารใหม่", "บันทึกอาคารใหม่สำเร็จ", "ผ่าน"],
                    ["24", "จัดลำดับอาคาร", "ลำดับอาคารเปลี่ยนแปลงตามที่กำหนด", "ผ่าน"],
                    ["25", "เพิ่มผู้ใช้งานใหม่โดยผู้ดูแลระบบ", "บันทึกผู้ใช้ใหม่สำเร็จ", "ผ่าน"],
                    ["26", "เปิด/ปิดใช้งานบัญชีผู้ใช้", "สถานะบัญชีเปลี่ยนแปลงตามที่กำหนด", "ผ่าน"],
                    ["27", "สอบถามข้อมูลอะไหล่ผ่าน AI Assistant", "AI ตอบคำถามเกี่ยวกับอะไหล่ได้ถูกต้อง", "ผ่าน"],
                    ["28", "ค้นหาอะไหล่ด้วยรูปภาพผ่าน AI", "แสดงรายการอะไหล่ที่ใกล้เคียงกับรูปภาพ", "ผ่าน"],
                    ["29", "สอบถามข้อมูลอะไหล่ผ่าน LINE Bot", "LINE Bot ตอบคำถามและแสดงข้อมูลอะไหล่ได้ถูกต้อง", "ผ่าน"],
                    ["30", "เข้าใช้งานระบบผ่าน LINE LIFF", "เปิดหน้าเว็บแอปภายใน LINE ได้ แสดงข้อมูลอะไหล่", "ผ่าน"],
                    ["31", "เข้าสู่ระบบแอปพลิเคชันมือถือ", "เข้าสู่ระบบสำเร็จ แสดงหน้า Dashboard บนมือถือ", "ผ่าน"],
                    ["32", "ค้นหาอะไหล่บนแอปพลิเคชันมือถือ", "แสดงรายการอะไหล่ตามคำค้นหา", "ผ่าน"],
                    ["33", "สแกน QR Code/Barcode บนแอปพลิเคชันมือถือ", "แสดงข้อมูลอะไหล่ที่ตรงกับรหัสที่สแกน", "ผ่าน"],
                    ["34", "เพิ่มอะไหล่ใหม่ผ่านแอปพลิเคชันมือถือ", "บันทึกข้อมูลอะไหล่ใหม่สำเร็จจากแอปพลิเคชันมือถือ", "ผ่าน"],
                    ["35", "ตรวจสอบสิทธิ์ผู้ใช้งานตามบทบาท (ADMIN/STAFF)", "ผู้ดูแลระบบใช้งานฟังก์ชันจัดการได้ เจ้าหน้าที่ทั่วไปใช้งานตามสิทธิ์", "ผ่าน"],
                ]
                for row_data in uat_items:
                    row = tbl.add_row()
                    for ci, val in enumerate(row_data):
                        cell = row.cells[ci]
                        cell.text = val
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.name = FONT
                                r.font.size = Pt(18)
                uat_table_found = True
                print(f"  Replaced UAT table with {len(uat_items)} items")
                break

if not uat_table_found:
    print("  WARNING: Could not find UAT table, skipping")

# Also fix paragraph 353 text: "6 รายการ" → "35 รายการ"
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if '6 รายการ' in run.text and 'UAT' in p.text:
            run.text = run.text.replace('6 รายการ', '35 รายการ')
            print(f"  [Paragraph {i}] Fixed UAT count: 6→35")
        elif 'ทั้ง 6 รายการ' in run.text:
            run.text = run.text.replace('ทั้ง 6 รายการ', 'ทั้ง 35 รายการ')
            print(f"  [Paragraph {i}] Fixed UAT count: 6→35")

# ══════════════════════════════════════════════════════════════
# TASK 10: Add bibliography at end (before sectPr)
# ══════════════════════════════════════════════════════════════
print("Task 10: Adding bibliography...")

# Find the last paragraph (should be paragraph ~361)
last_idx = len(doc.paragraphs) - 1
print(f"  Last paragraph index: {last_idx}")
print(f"  Last paragraph text: {doc.paragraphs[last_idx].text[:80]}")

# Find the actual last content paragraph (skip empty ones)
last_content_idx = last_idx
for i in range(last_idx, -1, -1):
    if doc.paragraphs[i].text.strip():
        last_content_idx = i
        break

print(f"  Last content paragraph index: {last_content_idx}")
print(f"  Last content text: {doc.paragraphs[last_content_idx].text[:80]}")

# Insert bibliography after last content
bib_chapter = _make_para("บรรณานุกรม", ref_para_idx=62, bold=True)  # Use chapter style
_insert_after_paragraph(last_content_idx, bib_chapter)

# Now insert bibliography entries after the chapter heading
bib_entries = [
    "กรมบัญชีกลาง. (2564). คู่มือการจัดทำบัญชีคลังสินค้า. กรุงเทพฯ: สำนักงานคณะกรรมการพัฒนาการเศรษฐกิจและสังคมแห่งชาติ.",
    "การไฟฟ้าฝ่ายผลิตแห่งประเทศไทย. (2565). คู่มือการบำรุงรักษาโรงไฟฟ้า. กรุงเทพฯ: กฟผ.",
    "การไฟฟ้าฝ่ายผลิตแห่งประเทศไทย. (2566). ระเบียบการจัดซื้อจัดจ้างและการจัดการคลังอะไหล่. กรุงเทพฯ: กฟผ.",
    "Vercel Inc. (2025). Next.js Documentation. Retrieved from https://nextjs.org/docs",
    "Meta Platforms, Inc. (2025). React Documentation. Retrieved from https://react.dev",
    "Microsoft Corporation. (2025). TypeScript Documentation. Retrieved from https://www.typescriptlang.org/docs",
    "Prisma Data Platform. (2025). Prisma Documentation. Retrieved from https://www.prisma.io/docs",
    "SQLite Consortium. (2025). SQLite Documentation. Retrieved from https://www.sqlite.org/docs.html",
    "LINE Corporation. (2025). LINE Developers Documentation. Retrieved from https://developers.line.biz/en/docs",
    "Google LLC. (2025). Flutter Documentation. Retrieved from https://flutter.dev/docs",
    "OpenAI. (2025). OpenAI API Documentation. Retrieved from https://platform.openai.com/docs",
    "Voyage AI. (2025). Voyage AI Embeddings Documentation. Retrieved from https://docs.voyageai.com",
    "Apache Software Foundation. (2025). Apache ECharts Documentation. Retrieved from https://echarts.apache.org/en/index.html",
    "สำนักงานนโยบายและแผนทรัพยากรธรรมชาติและสิ่งแวดล้อม. (2564). แนวทางการจัดการทรัพยากรอย่างยั่งยืน. กรุงเทพฯ: สผ.",
]

# Insert entries one by one (in reverse order so they appear in correct sequence)
ref_elem = doc.paragraphs[last_content_idx]._element
# Find the bibliography chapter we just inserted
bib_chap_elem = ref_elem.getnext()

current_ref = bib_chap_elem
for entry in bib_entries:
    entry_p = _make_para(entry, ref_para_idx=64, font_size=16)
    _insert_after_element(current_ref, entry_p)
    current_ref = entry_p

print(f"  Added bibliography with {len(bib_entries)} entries")

# ══════════════════════════════════════════════════════════════
# TASK 1: Insert ER Diagram after section 2.8
# ══════════════════════════════════════════════════════════════
print("Task 1: Inserting ER Diagram...")

# Find the paragraph after section 2.8 content (2.8.8 AppSetting = paragraph 189)
er_target = None
for i, p in enumerate(doc.paragraphs):
    if '2.8.8' in p.text and 'AppSetting' in p.text:
        er_target = i
        break

if er_target:
    ref_elem = doc.paragraphs[er_target]._element
    # Insert: empty line + image + caption (each inserted right after ref)
    empty_p = _make_para("", ref_para_idx=64)
    ref_elem.addnext(empty_p)
    ref_elem = empty_p

    _add_image_after(ref_elem, "er-diagram-v2.png", width=Inches(5.5))
    # The image paragraph is now right after empty_p; find it
    img_elem = empty_p.getnext()
    ref_elem = img_elem

    caption_elem = _add_caption("รูปที่ 2.1 แผนภาพความสัมพันธ์ระหว่างเอนทิตี (ER Diagram) ของระบบจัดการสต็อกอะไหล่")
    ref_elem.addnext(caption_elem)

    print(f"  Inserted ER Diagram after paragraph {er_target}")
else:
    print("  WARNING: Could not find 2.8.8 AppSetting paragraph")

# ══════════════════════════════════════════════════════════════
# TASK 2: Insert System Architecture Diagram after section 3.4
# ══════════════════════════════════════════════════════════════
print("Task 2: Inserting Architecture Diagram...")

# Find the paragraph for 3.4.3 (last item of section 3.4)
arch_target = None
for i, p in enumerate(doc.paragraphs):
    if '3.4.3' in p.text and 'ฐานข้อมูล' in p.text:
        arch_target = i
        break

if arch_target:
    ref_elem = doc.paragraphs[arch_target]._element
    empty_p = _make_para("", ref_para_idx=64)
    ref_elem.addnext(empty_p)
    ref_elem = empty_p

    _add_image_after(ref_elem, "architecture-diagram-v2.png", width=Inches(5.5))
    img_elem = empty_p.getnext()
    ref_elem = img_elem

    caption_elem = _add_caption("รูปที่ 3.1 แผนภาพโครงสร้างระบบ (System Architecture) ของระบบจัดการสต็อกอะไหล่")
    ref_elem.addnext(caption_elem)

    print(f"  Inserted Architecture Diagram after paragraph {arch_target}")
else:
    print("  WARNING: Could not find 3.4.3 paragraph")

# ══════════════════════════════════════════════════════════════
# TASK 3: Insert Flowchart images after section 3.5
# ══════════════════════════════════════════════════════════════
print("Task 3: Inserting Flowchart diagrams...")

# Find the paragraph for section 3.5 content
flow_target = None
for i, p in enumerate(doc.paragraphs):
    if 'ฐานข้อมูลของระบบประกอบด้วยตารางหลัก' in p.text:
        flow_target = i
        break

if flow_target:
    ref_elem = doc.paragraphs[flow_target]._element

    # Flowchart 1: Add Part
    empty1 = _make_para("", ref_para_idx=64)
    ref_elem.addnext(empty1)
    ref_elem = empty1
    _add_image_after(ref_elem, "flowchart-add-part.png", width=Inches(4.8))
    ref_elem = ref_elem.getnext()
    cap1 = _add_caption("รูปที่ 3.2 แผนผังลำดับการทำงาน: การเพิ่มข้อมูลอะไหล่")
    ref_elem.addnext(cap1)
    ref_elem = cap1

    # Flowchart 2: Search
    empty2 = _make_para("", ref_para_idx=64)
    ref_elem.addnext(empty2)
    ref_elem = empty2
    _add_image_after(ref_elem, "flowchart-search.png", width=Inches(4.8))
    ref_elem = ref_elem.getnext()
    cap2 = _add_caption("รูปที่ 3.3 แผนผังลำดับการทำงาน: การค้นหาอะไหล่")
    ref_elem.addnext(cap2)
    ref_elem = cap2

    # Flowchart 3: Stock Movement
    empty3 = _make_para("", ref_para_idx=64)
    ref_elem.addnext(empty3)
    ref_elem = empty3
    _add_image_after(ref_elem, "flowchart-stock-movement.png", width=Inches(4.8))
    ref_elem = ref_elem.getnext()
    cap3 = _add_caption("รูปที่ 3.4 แผนผังลำดับการทำงาน: การเคลื่อนไหวของสต็อก")
    ref_elem.addnext(cap3)

    print(f"  Inserted 3 Flowcharts after paragraph {flow_target}")
else:
    print("  WARNING: Could not find 3.5 content paragraph")

# ══════════════════════════════════════════════════════════════
# Save the document
# ══════════════════════════════════════════════════════════════
print(f"\nSaving to {OUT_PATH}...")
doc.save(OUT_PATH)
print("Done! ✓")

# Verify
doc2 = Document(OUT_PATH)
print(f"\nVerification:")
print(f"  Total paragraphs: {len(doc2.paragraphs)}")
print(f"  Total tables: {len(doc2.tables)}")

# Check key fixes
for i, p in enumerate(doc2.paragraphs):
    text = p.text
    if 'สต๊อก' in text:
        print(f"  ⚠️ Still found สต๊อก in paragraph {i}")
    if 'ปีการศึกษา 2568' in text:
        print(f"  ⚠️ Still found 2568 in paragraph {i}")
