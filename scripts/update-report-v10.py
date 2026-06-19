"""Build report-v10b-final.docx by cloning existing paragraph XML templates
from report-v9-final.docx so the inserted text matches the original format
(font TH SarabunPSK sz 36, bold headings, List Paragraph bullets, thaiDistribute).

Approach (cloning, not hand-built):
  - Heading template: clone para 468 (3.11 heading), set bold run text.
  - Body template:    clone para 469 (3.11 body, thaiDistribute justify).
  - Bullet template:  clone para 460 (List Paragraph bullet, numId 16).

Then insert the cloned paragraphs after the existing 3.11 body paragraph,
keeping only the 3.11 heading and replacing/augmenting its body with three
subsections (3.11.1 / 3.11.2 / 3.11.3).
"""
import sys
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

SRC = 'docs/report-v9-final.docx'
DST = 'docs/report-v10b-final.docx'

doc = Document(SRC)

# ---- locate template + anchor paragraphs in the BODY (not TOC) ----
def find_para(pred):
    for p in doc.paragraphs:
        if pred(p):
            return p
    return None

heading_template = None   # 3.11 heading (bold, sz36)
body_template = None      # 3.11 body paragraph (thaiDistribute)
bullet_template = None    # a List Paragraph bullet
anchor = None             # the 3.11 body paragraph we insert after

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == '3.11 สรุปขั้นตอนการดำเนินงาน' and '....' not in t:
        heading_template = p
        # next non-empty paragraph is the body
        for q in doc.paragraphs[i+1:]:
            if q.text.strip():
                body_template = q
                anchor = q
                break
    if p.style.name == 'List Paragraph' and t and '....' not in t and bullet_template is None:
        bullet_template = p

print(f"heading_template para: text={heading_template.text.strip()[:40]!r}")
print(f"body_template para:    text={body_template.text.strip()[:40]!r}")
print(f"bullet_template para:  text={bullet_template.text.strip()[:40]!r} style={bullet_template.style.name}")
assert heading_template and body_template and bullet_template and anchor

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def set_run_text(r_elem, text):
    """Replace text of the first w:t in a run, removing any others."""
    t_elems = r_elem.findall(qn('w:t'))
    if not t_elems:
        t = r_elem.makeelement(qn('w:t'), {})
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r_elem.append(t)
    else:
        t_elems[0].text = text
        t_elems[0].set(qn('xml:space'), 'preserve')
        for extra in t_elems[1:]:
            r_elem.remove(extra)

def strip_lastRenderedPageBreak(p_elem):
    for r in p_elem.findall(qn('w:r')):
        for lpb in r.findall(qn('w:lastRenderedPageBreak')):
            r.remove(lpb)

def make_from_template(template_para, text, *, is_heading=False):
    """Clone a template paragraph element and set its text.
    For headings we keep/ensure bold; for body we keep the existing run props.
    """
    new_elem = deepcopy(template_para._element)
    strip_lastRenderedPageBreak(new_elem)
    # remove w14:paraId/textId to avoid duplicates
    for attr in (qn('w14:paraId'), qn('w14:textId')):
        if new_elem.get(attr) is not None:
            del new_elem.attrib[attr]

    runs = new_elem.findall(qn('w:r'))
    if not runs:
        # build a minimal run
        r = new_elem.makeelement(qn('w:r'), {})
        new_elem.append(r)
        runs = [r]

    # keep only the first run, drop the rest
    for extra in runs[1:]:
        new_elem.remove(extra)
    first_run = runs[0]

    # ensure rPr exists with correct font/size
    rPr = first_run.find(qn('w:rPr'))
    if rPr is None:
        rPr = first_run.makeelement(qn('w:rPr'), {})
        first_run.insert(0, rPr)

    if is_heading:
        # ensure bold
        if rPr.find(qn('w:b')) is None:
            rPr.append(rPr.makeelement(qn('w:b'), {}))
        if rPr.find(qn('w:bCs')) is None:
            rPr.append(rPr.makeelement(qn('w:bCs'), {}))

    set_run_text(first_run, text)
    return new_elem

# ---- content to insert ----
# (kind, text)  kind in {heading, body, bullet}
CONTENT = [
    ('heading', '3.11.1 สรุปผลการดำเนินงาน'),
    ('body', 'จากการดำเนินงานทั้งหมด ระบบจัดการสต็อกอะไหล่ได้ถูกพัฒนาขึ้นเพื่อแก้ไขปัญหาการจัดเก็บและตรวจสอบข้อมูลอะไหล่ โดยระบบสามารถจัดการข้อมูลอะไหล่ แสดงสถานะสต็อก บันทึกประวัติการเคลื่อนไหว นำเข้าข้อมูลจาก Excel สร้าง QR Code/Barcode รองรับ AI Assistant พัฒนาแอปพลิเคชันมือถือด้วย Flutter และเชื่อมต่อกับ LINE Bot เพื่ออำนวยความสะดวกในการสอบถามข้อมูลสต็อก'),
    ('body', 'จากการนำระบบไปใช้งานเบื้องต้น พบว่าระบบสามารถลดระยะเวลาในการค้นหาอะไหล่จากเฉลี่ย 7.5 นาที เหลือไม่เกิน 1 นาที ลดเวลาการบันทึกข้อมูลอะไหล่ใหม่จากหลายขั้นตอนเหลือเพียงการถ่ายภาพและตรวจสอบข้อมูลที่ AI แนะนำ และสามารถติดตามประวัติการเคลื่อนไหวของสต็อกได้แบบเรียลไทม์ ส่งผลให้การจัดการอะไหล่มีประสิทธิภาพมากขึ้น นอกจากนี้ ผลการทดสอบการยอมรับระบบ (User Acceptance Test) จำนวน 35 รายการผ่านเกณฑ์ทุกข้อ แสดงว่าระบบสามารถทำงานได้ตามวัตถุประสงค์ที่กำหนดไว้'),

    ('heading', '3.11.2 ข้อจำกัดของระบบ'),
    ('body', 'ถึงแม้ระบบจะสามารถทำงานได้ตามวัตถุประสงค์หลัก แต่ยังคงมีข้อจำกัดบางประการที่ควรคำนึงถึงในการพัฒนาเพิ่มเติม ดังนี้'),
    ('bullet', 'ระบบยังไม่ได้เชื่อมต่อกับ SAP/ERP ขององค์กร ทำให้ข้อมูลการรับ จ่าย และคงเหลือของอะไหล่ต้องบันทึกผ่านระบบนี้เป็นหลัก ยังไม่สามารถดึงหรือส่งข้อมูลไปยังระบบ ERP ขององค์กรได้โดยอัตโนมัติ'),
    ('bullet', 'ระบบยังไม่ได้เชื่อมต่อกับระบบจัดซื้อ (Procurement) การขอซื้อหรือสั่งซื้ออะไหล่ยังคงต้องดำเนินการผ่านช่องทางเดิมแยกต่างหากนอกระบบนี้'),
    ('bullet', 'แอปพลิเคชันบนอุปกรณ์เคลื่อนที่ในปัจจุบันใช้เทคโนโลยี Flutter แบบ Web-based และ LINE LIFF ยังไม่ใช่ Native Mobile Application ที่สามารถติดตั้งผ่าน App Store หรือ Play Store โดยตรง'),
    ('bullet', 'ความแม่นยำของระบบ AI ไม่ว่าจะเป็นการค้นหาด้วยรูปภาพ การวิเคราะห์ข้อมูลอะไหล่จากภาพ หรือการตอบคำถามผ่านแชท ยังขึ้นกับคุณภาพของรูปภาพ แสงสว่าง องศาการถ่าย และความชัดเจนของข้อมูลที่ผู้ใช้ป้อนเข้า'),

    ('heading', '3.11.3 แนวทางการพัฒนาในอนาคต'),
    ('body', 'เพื่อให้ระบบสามารถตอบสนองความต้องการของหน่วยงานได้ครอบคลุมมากยิ่งขึ้นในอนาคต ควรพิจารณาแนวทางการพัฒนาดังต่อไปนี้'),
    ('bullet', 'เชื่อมต่อระบบกับ SAP/ERP ขององค์กร เพื่อซิงค์ข้อมูลอะไหล่ จำนวนคงเหลือ และมูลค่าต้นทุนแบบอัตโนมัติ ลดภาระการป้อนข้อมูลซ้ำซ้อนและเพิ่มความถูกต้องของข้อมูล'),
    ('bullet', 'เพิ่มระบบแจ้งเตือนอะไหล่ใกล้หมดสต็อก (Low Stock Notification) ผ่าน LINE หรืออีเมล ช่วยให้ฝ่ายจัดซื้อและผู้ดูแลคลังสินค้าสามารถเตรียมความพร้อมก่อนเกิดภาวะขาดแคลน'),
    ('bullet', 'พัฒนาโมดูล Predictive Spare Parts โดยใช้เทคโนโลยี AI วิเคราะห์แนวโน้มการใช้อะไหล่จากประวัติการเบิกจ่าย เพื่อคาดการณ์ความต้องการในอนาคตและช่วยวางแผนสต็อกล่วงหน้า'),
    ('bullet', 'เพิ่ม Dashboard สำหรับวิเคราะห์แนวโน้มการใช้ Spare Part ในรูปแบบกราฟและรายงาน เช่น รายการอะไหล่ที่เบิกบ่อยที่สุด ช่วงเวลาที่มีอัตราการใช้สูงสุด และการเปรียบเทียบยอดการใช้ระหว่างเดือน'),
    ('bullet', 'ศึกษาและประยุกต์ใช้เทคโนโลยี RFID สำหรับติดตามอะไหล่เข้า–ออกคลังสินค้า ซึ่งจะช่วยเพิ่มความรวดเร็วในการตรวจนับสต็อกและลดความผิดพลาดจากการสแกนได้มากกว่า QR Code และ Barcode'),
    ('bullet', 'พัฒนา Native Mobile Application บน iOS และ Android เพื่อรองรับการทำงานแบบ Offline-First สำหรับพนักงานที่ต้องปฏิบัติงานในพื้นที่ที่สัญญาณอินเทอร์เน็ตไม่เสถียร'),
]

# ---- insert cloned paragraphs after anchor (the existing 3.11 body) ----
parent = anchor._element.getparent()
idx = list(parent).index(anchor._element)

template_for = {
    'heading': heading_template,
    'body': body_template,
    'bullet': bullet_template,
}

for kind, text in CONTENT:
    new_elem = make_from_template(template_for[kind], text, is_heading=(kind == 'heading'))
    idx += 1
    parent.insert(idx, new_elem)

print(f"Inserted {len(CONTENT)} paragraphs")

doc.save(DST)
print(f"Saved: {DST}")
