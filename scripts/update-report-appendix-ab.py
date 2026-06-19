"""Rebuild Appendix A (Project Structure) as grouped headings + bullets, and
append 4 more code blocks to Appendix B (stock.ts, AI orchestrator, LINE
webhook + LIFF auth, part-text-search) — cloning existing paragraph templates
so formatting matches the v9-original style exactly.
"""
import sys
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

SRC = 'docs/report-v10b-final.docx'
DST = 'docs/report-v10b-final.docx'

doc = Document(SRC)
paras = doc.paragraphs

# ---------- locate templates ----------
# Bullet template: a List Paragraph bullet (from 3.10 area)
bullet_template = None
for p in paras:
    if p.style.name == 'List Paragraph' and p.text.strip() and '....' not in p.text:
        bullet_template = p
        break

# Subheading template: clone the appendix-B title paragraph ('โค้ดสำคัญ...')
# but we want a left-aligned bold subheading, so we use the 3.11.1 style we inserted.
subheading_template = None
for p in paras:
    if p.text.strip().startswith('3.11.1'):
        subheading_template = p
        break

# Code block template: the shaded Consolas paragraph holding Prisma Schema
code_template = None
code_template_idx = None
for i, p in enumerate(paras):
    pPr = p._p.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:shd')) is not None:
        code_template = p
        code_template_idx = i
        break

print(f"bullet_template: {bullet_template.text.strip()[:30]!r}")
print(f"subheading_template: {subheading_template.text.strip()[:30]!r}")
print(f"code_template: {code_template.text.strip()[:30]!r} (para {code_template_idx})")
assert bullet_template and subheading_template and code_template

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def set_run_text(r_elem, text):
    t_elems = r_elem.findall(qn('w:t'))
    if not t_elems:
        t = r_elem.makeelement(qn('w:t'), {})
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r_elem.append(t)
    else:
        t_elems[0].text = text
        t_elems[0].set(qn('xml:space'), 'preserve')
        for extra in t_elems[1:]:
            r_elem.remove(extra)

def strip_extras(p_elem):
    # remove w14:paraId/textId and lastRenderedPageBreak
    for attr in (qn('w14:paraId'), qn('w14:textId')):
        if p_elem.get(attr) is not None:
            del p_elem.attrib[attr]
    for r in p_elem.findall(qn('w:r')):
        for lpb in r.findall(qn('w:lastRenderedPageBreak')):
            r.remove(lpb)

def make_subheading(text):
    el = deepcopy(subheading_template._element)
    strip_extras(el)
    # remove center justification -> left
    pPr = el.find(qn('w:pPr'))
    jc = pPr.find(qn('w:jc'))
    if jc is not None:
        pPr.remove(jc)
    runs = el.findall(qn('w:r'))
    for extra in runs[1:]:
        el.remove(extra)
    set_run_text(runs[0], text)
    return el

def make_bullet(text):
    el = deepcopy(bullet_template._element)
    strip_extras(el)
    runs = el.findall(qn('w:r'))
    for extra in runs[1:]:
        el.remove(extra)
    set_run_text(runs[0], text)
    return el

def make_body(text):
    """A normal body paragraph (thaiDistribute, sz36). Clone 3.11.1 first body."""
    # use body_template = the paragraph right after 3.11.1 heading
    el = deepcopy(body_template._element)
    strip_extras(el)
    runs = el.findall(qn('w:r'))
    for extra in runs[1:]:
        el.remove(extra)
    set_run_text(runs[0], text)
    return el

# body template = paragraph after 3.11.1 heading
body_template = None
for i, p in enumerate(paras):
    if p.text.strip().startswith('3.11.1'):
        body_template = paras[i + 1]
        break
print(f"body_template: {body_template.text.strip()[:30]!r}")

def make_code_block(lines):
    """Clone the shaded code paragraph; replace its text with multi-line via <w:br/>."""
    el = deepcopy(code_template._element)
    strip_extras(el)
    # Keep first run, drop the rest
    runs = el.findall(qn('w:r'))
    first_run = runs[0]
    for extra in runs[1:]:
        el.remove(extra)
    # Clear first run's children except rPr
    rPr = first_run.find(qn('w:rPr'))
    for child in list(first_run):
        if child is not rPr:
            first_run.remove(child)
    # Build text with line breaks
    for idx, line in enumerate(lines):
        if idx > 0:
            br = OxmlElement('w:br')
            first_run.append(br)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = line
        first_run.append(t)
    return el

# ============================================================
# PART 1 — rebuild Appendix A as grouped headings + bullets
# ============================================================
a_start = None
a_end = None
for i, p in enumerate(paras):
    t = p.text.strip()
    if t == 'ภาคผนวก ก' and a_start is None:
        a_start = i
    elif a_start is not None and t == 'ภาคผนวก ข':
        a_end = i
        break
print(f"\nAppendix A: {a_start} -> {a_end}")

# Keep heading (a_start) and title (a_start+1). Delete body paras a_start+2 .. a_end-1.
title_elem = paras[a_start + 1]._element
to_delete = [paras[i]._element for i in range(a_start + 2, a_end)]
for el in to_delete:
    el.getparent().remove(el)
print(f"Deleted {len(to_delete)} old appendix-A body paragraphs")

# New grouped content: (kind, text)
APPENDIX_A = [
    ('sub', 'ส่วนหน้าเว็บ (src/app)'),
    ('bullet', 'dashboard/ — หน้าสรุป Dashboard แสดงจำนวนอะไหล่ สต็อก และประวัติการเคลื่อนไหวล่าสุด'),
    ('bullet', 'parts/ — จัดการอะไหล่ (เพิ่มใหม่ รายละเอียด และรายการทั้งหมด)'),
    ('bullet', 'movements/ — ประวัติการรับ–จ่าย–ปรับสต็อก (Stock Movement)'),
    ('bullet', 'scan/ — สแกน QR Code และ Barcode เพื่อค้นหาอะไหล่'),
    ('bullet', 'import/ — นำเข้าข้อมูลอะไหล่จากไฟล์ Excel'),
    ('bullet', 'categories/, buildings/, blocks/ — จัดการหมวดหมู่ อาคาร และบล็อก'),
    ('bullet', 'users/ — จัดการผู้ใช้งาน (เฉพาะผู้ดูแลระบบ)'),
    ('bullet', 'assistant/ — หน้าสนทนากับ AI Assistant'),
    ('bullet', 'settings/ — ตั้งค่าระบบและเปลี่ยนรหัสผ่าน'),
    ('sub', 'ส่วน API (src/app/api)'),
    ('bullet', 'auth/ — ระบบล็อกอินและจัดการเซสชัน'),
    ('bullet', 'admin/ — API สำหรับผู้ดูแลระบบ (จัดการผู้ใช้ ล้างฐานข้อมูล)'),
    ('bullet', 'ai/ — API สำหรับสนทนาและดำเนินการคำสั่งของ AI'),
    ('bullet', 'liff/ — API สำหรับ LINE LIFF mini app'),
    ('bullet', 'line/ — LINE Webhook รับข้อความและภาพจากแชท LINE'),
    ('bullet', 'mobile/ — REST API สำหรับแอปพลิเคชันมือถือ (Flutter)'),
    ('bullet', 'parts/, movements/, categories/, buildings/, blocks/ — API ของแต่ละโมดูล'),
    ('bullet', 'import/, export/ — นำเข้าและส่งออกข้อมูล'),
    ('bullet', 'dashboard/ — API สรุปข้อมูลสำหรับหน้า Dashboard'),
    ('bullet', 'cron/ — Cron jobs สำหรับงานที่กำหนดเวลา'),
    ('sub', 'ส่วน LIFF (src/app/liff)'),
    ('bullet', 'add-part/ — เพิ่มอะไหล่ผ่าน LINE LIFF พร้อม AI แนะนำข้อมูลจากภาพ'),
    ('bullet', 'search/ — ค้นหาอะไหล่ผ่าน LINE LIFF (รวมค้นหาด้วยรูปภาพ)'),
    ('bullet', 'stock-move/ — รับ/จ่ายสต็อกผ่าน LINE LIFF'),
    ('sub', 'ส่วน Core (src/lib)'),
    ('bullet', 'stock.ts — โลจิกการเคลื่อนไหวสต็อก ตรวจจำนวนติดลบ และบันทึกประวัติใน transaction'),
    ('bullet', 'ai-assistant/ — เอนจิน AI Assistant (orchestrator, tools, pending actions)'),
    ('bullet', 'line-chat/ — การจัดการบทสนทนาและ Flex Message ของ LINE Bot'),
    ('bullet', 'part-lookup.ts — บริการค้นหาอะไหล่จากรหัส (QR/Barcode/Part Number)'),
    ('bullet', 'part-text-search.ts — ค้นหาอะไหล่ด้วย Vector Embedding (ความหมายใกล้เคียง)'),
    ('bullet', 'embeddings.ts — สร้างและเปรียบเทียบ Vector Embedding ของภาพและข้อความ'),
    ('bullet', 'excel.ts — อ่าน/เขียนไฟล์ Excel'),
    ('bullet', 'auth.ts, session.ts — ตัวช่วยตรวจสอบสิทธิ์และจัดการเซสชัน'),
    ('bullet', 'prisma.ts — Prisma Client singleton'),
    ('sub', 'ส่วน Component และคอนฟิก (src/components, root)'),
    ('bullet', 'components/layout/ — Sidebar, Header, BottomNav'),
    ('bullet', 'components/ui/ — ชุดคอมโพเนนต์ shadcn/ui (Button, Table, Dialog, ...)'),
    ('bullet', 'prisma/schema.prisma — Prisma Schema ของฐานข้อมูล'),
    ('bullet', '.env.example — ตัวอย่าง environment variables'),
    ('bullet', 'next.config.ts, package.json, tsconfig.json — คอนฟิก Next.js และ TypeScript'),
]

ref = title_elem
for kind, text in APPENDIX_A:
    if kind == 'sub':
        el = make_subheading(text)
    else:
        el = make_bullet(text)
    ref.addnext(el)
    ref = el
print(f"Inserted {len(APPENDIX_A)} appendix-A paragraphs")

# ============================================================
# PART 2 — append 4 code blocks to Appendix B
# ============================================================
# Re-fetch paragraphs (indices shifted)
paras = doc.paragraphs
b_title_idx = None
pagebreak_idx = None
for i, p in enumerate(paras):
    if p.text.strip().startswith('โค้ดสำคัญของระบบ — Prisma Schema'):
        b_title_idx = i
        # the Prisma code block is next; the page break para follows it
        break
print(f"\nAppendix B title at para {b_title_idx}")

# find the page-break paragraph that ends appendix B (the one with w:br type=page)
pagebreak_elem = None
for i in range(b_title_idx, len(paras)):
    p = paras[i]
    for r in p._element.findall(qn('w:r')):
        for br in r.findall(qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                pagebreak_elem = p._element
                break
    if pagebreak_elem is not None:
        break
print(f"page-break para found: {pagebreak_elem is not None}")

# We'll insert new subheadings + code blocks BEFORE the page-break paragraph.
# Insertion point = element before pagebreak (i.e. after the last code line of Prisma).
# Simplest: insert right before pagebreak_elem.
def insert_before(ref_elem, new_elem):
    parent = ref_elem.getparent()
    index = list(parent).index(ref_elem)
    parent.insert(index, new_elem)

# Code snippets (trimmed to the essential parts)
SNIPPETS = [
    (
        'โค้ดสำคัญของระบบ — Stock Movement (src/lib/stock.ts)',
        [
            'export class StockError extends Error {',
            '  constructor(message: "PART_NOT_FOUND" | "INSUFFICIENT_STOCK" |',
            '    "NEGATIVE_STOCK" | "CONCURRENT_MODIFICATION") {',
            '    super(message);',
            '    this.name = "StockError";',
            '  }',
            '}',
            '',
            'async function createStockMovementWithClient(db, params) {',
            '  const { partId, userId, type, quantity, note } = params;',
            '  const part = await db.part.findUnique({ where: { id: partId } });',
            '  if (!part) throw new StockError("PART_NOT_FOUND");',
            '',
            '  const quantityBefore = part.quantity;',
            '  let quantityAfter = part.quantity;',
            '  if (type === "STOCK_IN") quantityAfter = part.quantity + quantity;',
            '  else if (type === "STOCK_OUT") quantityAfter = part.quantity - quantity;',
            '  else if (type === "ADJUSTMENT") quantityAfter = quantity;',
            '',
            '  if (quantityAfter < 0) {',
            '    throw type === "STOCK_OUT"',
            '      ? new StockError("INSUFFICIENT_STOCK")',
            '      : new StockError("NEGATIVE_STOCK");',
            '  }',
            '',
            '  // Atomic guarded update — refuses to write if quantity changed since read.',
            '  const updated = await db.part.updateMany({',
            '    where: { id: partId, quantity: quantityBefore },',
            '    data: { quantity: quantityAfter },',
            '  });',
            '  if (updated.count === 0) throw new StockError("CONCURRENT_MODIFICATION");',
            '',
            '  const movement = await db.stockMovement.create({',
            '    data: { partId, userId, type, quantityBefore, quantityAfter, note },',
            '  });',
            '  return movement;',
            '}',
            '',
            'export async function createStockMovement(params, db?) {',
            '  if (db) return createStockMovementWithClient(db, params);',
            '  return prisma.$transaction(async (tx) => {',
            '    return createStockMovementWithClient(tx, params);',
            '  });',
            '}',
        ],
    ),
    (
        'โค้ดสำคัญของระบบ — AI Assistant Orchestrator (src/lib/ai-assistant/orchestrator.ts)',
        [
            'export async function runAiAssistant(input): Promise<AiAssistantResult> {',
            '  // Deterministic pre-router: skip LLM for common stock query patterns.',
            '  const directTool = tryDirectToolRouting(input.message);',
            '  if (directTool) {',
            '    const conversationId = await resolveConversationId(input);',
            '    if (conversationId && !input.skipSaveUserMessage) {',
            '      await saveMessage(conversationId, "user", input.message, "text");',
            '    }',
            '    const toolResult = await executeAiTool(directTool.name, directTool.args);',
            '    // Use LLM only for generating the natural-language summary',
            '    let reply = await generateReplyFromToolResult(toolResult.content, directTool.name);',
            '    reply = sanitizeWebReply(reply);',
            '    await saveMessage(conversationId, "assistant", reply, "text");',
            '    return { reply, conversationId, toolCalls: [{ name: directTool.name }] };',
            '  }',
            '  return runAiAssistantViaLlm(input);',
            '}',
            '',
            '// Deterministic pre-router: detect common stock query patterns',
            '// and return the tool call to execute directly, bypassing the LLM.',
            'function tryDirectToolRouting(message: string) {',
            '  const text = message.trim();',
            '  if (!text || text.length < 4) return null;',
            '  // "สถานะ/สรุป/เหลือเท่าไหร่ [keyword]" -> get_stock_summary',
            '  if (hasSummaryTerms(text) && hasInventoryContent(text)) {',
            '    const filters = extractInventoryFilters(text);',
            '    if (filters.keyword) {',
            '      return { name: "get_stock_summary", args: filters };',
            '    }',
            '  }',
            '  return null; // let LLM decide',
            '}',
        ],
    ),
    (
        'โค้ดสำคัญของระบบ — LINE Webhook (src/app/api/line/webhook/route.ts)',
        [
            'export async function POST(request: Request) {',
            '  const body = await request.text();',
            '  const signature = request.headers.get("x-line-signature") || "";',
            '  if (!signature) {',
            '    return NextResponse.json({ message: "Missing signature" }, { status: 401 });',
            '  }',
            '  if (!verifyLineSignature(body, signature)) {',
            '    return NextResponse.json({ message: "Invalid signature" }, { status: 401 });',
            '  }',
            '',
            '  const parsed = JSON.parse(body) as LineWebhookBody;',
            '  const events = parsed.events || [];',
            '  for (const event of events) {',
            '    if (event.type === "follow") { /* ส่งเมนูต้อนรับ */ continue; }',
            '    if (event.type === "postback" && event.postback?.data) {',
            '      const action = parseAction(event.postback.data);',
            '      const user = await findLineLinkedUser(event.source?.userId);',
            '      await handlePostback(action, user, event.replyToken, ...);',
            '      continue;',
            '    }',
            '    if (event.type !== "message") continue;',
            '    // ข้อความ/ภาพ -> ส่งให้ orchestrator วิเคราะห์และตอบ',
            '    const result = await orchestrate(userId, lineGroupId, text, isGroup, ...);',
            '    await sendLineReply(event.replyToken, buildReplyMessages(result));',
            '  }',
            '  return NextResponse.json({ ok: true });',
            '}',
        ],
    ),
    (
        'โค้ดสำคัญของระบบ — LIFF Authentication (src/lib/liff-auth.tsx)',
        [
            'async function getLiffProfile(): Promise<{ idToken: string }> {',
            '  const mod = await import("@line/liff");',
            '  const liff = mod.default;',
            '  await liff.init({ liffId: LIFF_ID, withLoginOnExternalBrowser: true });',
            '',
            '  if (!liff.isLoggedIn()) {',
            '    // จด path ปัจจุบันไว้ แล้วเปลี่ยนเส้นทางกลับหลังล็อกอิน LINE',
            '    localStorage.setItem("liff_login_redirect",',
            '      window.location.pathname + window.location.search);',
            '    liff.login({ redirectUri: window.location.href });',
            '    return new Promise(() => {});',
            '  }',
            '  const idToken = liff.getIDToken();',
            '  return { idToken };',
            '}',
            '',
            'async function verifyToken(idToken: string) {',
            '  const res = await fetch("/api/line/auth/verify-token", {',
            '    method: "POST",',
            '    headers: { "Content-Type": "application/json", "X-API-Key": API_KEY },',
            '    body: JSON.stringify({ idToken }),',
            '  });',
            '  return res.json();',
            '}',
        ],
    ),
    (
        'โค้ดสำคัญของระบบ — Hybrid Part Search (src/lib/part-text-search.ts)',
        [
            'export async function searchPartsByTextEmbedding(queryText, options = {}) {',
            '  const { minSimilarity = 0.35, limit = 10, plant, buildingId } = options;',
            '  if (!queryText.trim()) return [];',
            '',
            '  // แปลงข้อความค้นหาเป็น vector embedding',
            '  const queryVector = await embedText(queryText, "query");',
            '',
            '  // โหลดอะไหล่ที่มี textEmbedding และ active',
            '  const parts = await prisma.part.findMany({',
            '    where: { isActive: true, textEmbedding: { not: null } },',
            '    select: { id: true, partNumber: true, partName: true,',
            '      quantity: true, textEmbedding: true,',
            '      category: { select: { name: true } } },',
            '    take: 2000,',
            '  });',
            '',
            '  // คำนวณความใกล้เคียง (cosine similarity) ของแต่ละอะไหล่',
            '  const matches = [];',
            '  for (const part of parts) {',
            '    const partVector = bytesToFloat32(part.textEmbedding);',
            '    const similarity = cosineSimilarity(queryVector, partVector);',
            '    if (similarity >= minSimilarity) {',
            '      matches.push({ ...part, similarity });',
            '    }',
            '  }',
            '  // เรียงตามความใกล้เคียงมากไปน้อย',
            '  matches.sort((a, b) => b.similarity - a.similarity);',
            '  return matches.slice(0, limit);',
            '}',
        ],
    ),
]

# Insert in order before the page-break. Each block: subheading + blank + code + blank
for title, lines in SNIPPETS:
    insert_before(pagebreak_elem, make_subheading(title))
    # blank line
    blank = deepcopy(body_template._element)
    strip_extras(blank)
    runs = blank.findall(qn('w:r'))
    for r in runs:
        blank.remove(r)
    insert_before(pagebreak_elem, blank)
    insert_before(pagebreak_elem, make_code_block(lines))
    insert_before(pagebreak_elem, deepcopy(blank))

print(f"Inserted {len(SNIPPETS)} new code blocks into Appendix B")

doc.save(DST)
print(f"\nSaved: {DST}")
