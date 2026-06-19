"""
Generate internship report v8 — add TOC, list of figures, list of tables,
and appendices to report-v7-final.docx.

Changes from v7→v8:
  1. Fix screenshot figure numbers: 3.1–3.19 → 3.5–3.23 (diagrams keep 3.1–3.4)
  2. Insert สารบัญ (Table of Contents) after หนังสือรับรอง
  3. Insert สารบัญรูป (List of Figures)
  4. Insert สารบัญตาราง (List of Tables)
  5. Insert ภาคผนวก ก ข ง จ after บรรณานุกรม
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import os, copy, re

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMG_DIR = os.path.join(BASE, "docs", "images")
INPUT = os.path.join(BASE, "docs", "report-v7-final.docx")
OUT_PATH = os.path.join(BASE, "docs", "report-v8-final.docx")

FONT = 'TH SarabunPSK'

doc = Document(INPUT)
body = doc.element.body

# ── Reference paragraph indices ──
# Find useful reference paragraphs for cloning
REF_CHAPTER = None   # e.g. "บทที่ 1 บทนำ"
REF_SECTION = None   # e.g. "1.1 ความเป็นมา..."
REF_BODY = None      # normal body text
REF_CAPTION = None   # italic caption like "รูปที่ ..."

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if REF_CHAPTER is None and t.startswith('บทที่'):
        REF_CHAPTER = i
    if REF_SECTION is None and re.match(r'^\d+\.\d+\s', t):
        REF_SECTION = i
    if REF_BODY is None and len(t) > 40 and not t.startswith('บทที่') and not t.startswith('รูปที่'):
        REF_BODY = i
    if REF_CAPTION is None and t.startswith('รูปที่'):
        REF_CAPTION = i
    if all([REF_CHAPTER, REF_SECTION, REF_BODY, REF_CAPTION]):
        break

print(f"Reference paragraphs: chapter={REF_CHAPTER}, section={REF_SECTION}, body={REF_BODY}, caption={REF_CAPTION}")

# ══════════════════════════════════════════════════════════════
# Helper functions
# ══════════════════════════════════════════════════════════════
def _escape_xml(text):
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')

def _make_para(text, ref_idx=None, bold=False, alignment=None, italic=False, font_size=None):
    if ref_idx is None:
        ref_idx = REF_BODY or 64
    ref_elem = doc.paragraphs[ref_idx]._element
    new_p = copy.deepcopy(ref_elem)
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    ref_runs = ref_elem.findall(qn('w:r'))
    new_r = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
    if ref_runs:
        ref_rPr = ref_runs[0].find(qn('w:rPr'))
        if ref_rPr is not None:
            new_rPr = copy.deepcopy(ref_rPr)
            if bold:
                for tag in ['w:b', 'w:bCs']:
                    e = new_rPr.find(qn(tag))
                    if e is None:
                        e = parse_xml(f'<{tag} {nsdecls("w")}/>')
                        new_rPr.append(e)
                    else:
                        e.set(qn('w:val'), '1')
            else:
                for tag in ['w:b', 'w:bCs']:
                    e = new_rPr.find(qn(tag))
                    if e is not None:
                        e.set(qn('w:val'), '0')
            if italic:
                for tag in ['w:i', 'w:iCs']:
                    e = new_rPr.find(qn(tag))
                    if e is None:
                        e = parse_xml(f'<{tag} {nsdecls("w")}/>')
                        new_rPr.append(e)
                    else:
                        e.set(qn('w:val'), '1')
            if font_size:
                for tag in ['w:sz', 'w:szCs']:
                    e = new_rPr.find(qn(tag))
                    if e is not None:
                        e.set(qn('w:val'), str(font_size * 2))
            new_r.append(new_rPr)
    t = parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve">{_escape_xml(text)}</w:t>')
    new_r.append(t)
    new_p.append(new_r)
    if alignment:
        pPr = new_p.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            new_p.insert(0, pPr)
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="{alignment}"/>'))
        else:
            jc.set(qn('w:val'), alignment)
    return new_p

def _add_page_break_elem():
    return parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:r><w:br w:type="page"/></w:r>'
        f'</w:p>'
    )

def _add_image_after(ref_elem, filename, width=Inches(5.2)):
    path = os.path.join(IMG_DIR, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if os.path.exists(path):
        run.add_picture(path, width=width)
    else:
        run.text = f"[รูปภาพ: {filename}]"
    p_elem = p._element
    body.remove(p_elem)
    ref_elem.addnext(p_elem)
    return p_elem

def _make_table_elem(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = t._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = FONT
                r.font.size = Pt(18)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT
                    r.font.size = Pt(18)
    tbl_elem = t._tbl
    tbl_elem.getparent().remove(tbl_elem)
    return tbl_elem

def _dots(label, page=""):
    """Generate TOC line with dot leaders."""
    if page:
        # Approximate: fill with dots between label and page
        space = max(3, 55 - len(label) - len(page))
        return f"{label} {'.' * space} {page}"
    return label

# ══════════════════════════════════════════════════════════════
# STEP 1: Fix screenshot figure numbers 3.1–3.19 → 3.5–3.23
# ══════════════════════════════════════════════════════════════
print("Step 1: Fixing screenshot figure numbers...")

DIAGRAM_KEYWORDS = ['ER Diagram', 'System Architecture', 'แผนผังลำดับการทำงาน']
fix_count = 0

# Build a map of screenshot captions that need renumbering
# Diagrams keep 3.1-3.4, screenshots shift by +4
# We need to process ALL runs of a paragraph together, because the
# "รูปที่ 3.16" might be split across runs like "รูปที่ 3.1" + "6"

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t.startswith('รูปที่'):
        continue
    # Skip diagrams — they keep their numbers
    if any(k in t for k in DIAGRAM_KEYWORDS):
        continue
    # This is a screenshot caption — fix the number
    # Extract the full number from the combined text
    m = re.match(r'รูปที่\s*3\.(\d+)', t)
    if not m:
        continue
    old_num = int(m.group(1))
    new_num = old_num + 4
    old_str = f'3.{old_num}'
    new_str = f'3.{new_num}'

    # Find which run(s) contain the number and replace
    # Simple approach: just replace in any run that contains part of the number
    for run in p.runs:
        if old_str in run.text:
            run.text = run.text.replace(old_str, new_str)
            fix_count += 1
            print(f"  [{i}] {t[:50]} -> รูปที่ {new_str}")
            break
    else:
        # Number might be split across runs — try combining first two runs
        if len(p.runs) >= 2:
            combined = p.runs[0].text + p.runs[1].text
            if old_str in combined:
                # Put the full replacement in the first run, clear the second's number part
                p.runs[0].text = p.runs[0].text.replace(old_str, new_str)
                if p.runs[1].text.startswith(str(old_num)[-1:]):
                    p.runs[1].text = p.runs[1].text[len(str(old_num))-len(str(old_num).rstrip(str(old_num)[-1])):]
                fix_count += 1
                print(f"  [{i}] (split) {t[:50]} -> รูปที่ {new_str}")

print(f"  Fixed {fix_count} screenshot captions")

# ══════════════════════════════════════════════════════════════
# STEP 2: Collect all headings, figures, tables for TOC
# ══════════════════════════════════════════════════════════════
print("Step 2: Collecting TOC entries...")

toc_entries = []  # (level, text, page_placeholder)
fig_entries = []  # (fig_num, caption)
tbl_entries = []  # (tbl_num, caption)

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue

    # Chapters
    if t.startswith('บทที่'):
        toc_entries.append((0, t, ""))
    # Sections like 1.1, 2.3, 3.6
    elif re.match(r'^\d+\.\d+\s', t) and len(t) < 80:
        # Sub-sections like 1.3.1
        if re.match(r'^\d+\.\d+\.\d+\s', t):
            toc_entries.append((2, t, ""))
        else:
            toc_entries.append((1, t, ""))
    # Figure captions
    elif t.startswith('รูปที่'):
        fig_entries.append(t)
    # Standalone section headings like สารบัญ, บรรณานุกรม, ภาคผนวก
    elif t in ['สารบัญ', 'สารบัญรูป', 'สารบัญตาราง', 'บรรณานุกรม', 'ภาคผนวก']:
        toc_entries.append((0, t, ""))

# Table captions (manual — based on our knowledge of the document)
table_captions = [
    "ตารางที่ 1.1 ตัวชี้วัดประสิทธิภาพ (KPI) ก่อนและหลังการพัฒนาระบบ",
    "ตารางที่ 1.2 นิยามศัพท์เฉพาะ",
    "ตารางที่ 3.1 โครงสร้างตารางฐานข้อมูลของระบบ",
    "ตารางที่ 3.2 ผลการทดสอบการยอมรับระบบ (UAT)",
]

print(f"  Found {len(toc_entries)} TOC entries, {len(fig_entries)} figures, {len(table_captions)} tables")

# ══════════════════════════════════════════════════════════════
# STEP 3: Insert สารบัญ + สารบัญรูป + สารบัญตาราง
#         after หนังสือรับรอง (before บทที่ 1)
# ══════════════════════════════════════════════════════════════
print("Step 3: Inserting TOC pages...")

# Find the paragraph just before บทที่ 1
ch1_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('บทที่ 1'):
        ch1_idx = i
        break

if ch1_idx is None:
    print("  ERROR: Could not find บทที่ 1")
    sys.exit(1)

# Insert before บทที่ 1 by inserting after the paragraph just before it
# Find the last non-empty paragraph before บทที่ 1
pre_ch1 = ch1_idx - 1
while pre_ch1 > 0 and not doc.paragraphs[pre_ch1].text.strip():
    pre_ch1 -= 1

print(f"  บทที่ 1 at paragraph {ch1_idx}, inserting TOC after paragraph {pre_ch1}")

# Build all TOC elements, then insert in reverse order before บทที่ 1
toc_elems = []

# --- สารบัญ ---
toc_elems.append(_add_page_break_elem())
toc_title = _make_para("สารบัญ", ref_idx=REF_CHAPTER, bold=True, alignment='center')
toc_elems.append(toc_title)
toc_elems.append(_make_para("", ref_idx=REF_BODY))  # blank line

# Add TOC entries
for level, text, _ in toc_entries:
    indent = ""
    if level == 0:
        entry = _make_para(text, ref_idx=REF_SECTION, bold=True)
    elif level == 1:
        entry = _make_para(f"     {text}", ref_idx=REF_BODY)
    else:
        entry = _make_para(f"          {text}", ref_idx=REF_BODY, font_size=16)
    toc_elems.append(entry)

# --- สารบัญรูป ---
toc_elems.append(_add_page_break_elem())
fig_title = _make_para("สารบัญรูป", ref_idx=REF_CHAPTER, bold=True, alignment='center')
toc_elems.append(fig_title)
toc_elems.append(_make_para("", ref_idx=REF_BODY))

for fig_text in fig_entries:
    entry = _make_para(fig_text, ref_idx=REF_BODY)
    toc_elems.append(entry)

# --- สารบัญตาราง ---
toc_elems.append(_add_page_break_elem())
tbl_title = _make_para("สารบัญตาราง", ref_idx=REF_CHAPTER, bold=True, alignment='center')
toc_elems.append(tbl_title)
toc_elems.append(_make_para("", ref_idx=REF_BODY))

for tbl_cap in table_captions:
    entry = _make_para(tbl_cap, ref_idx=REF_BODY)
    toc_elems.append(entry)

# Insert all elements before บทที่ 1 (by inserting after pre_ch1 in reverse)
ref_elem = doc.paragraphs[pre_ch1]._element
for elem in reversed(toc_elems):
    ref_elem.addnext(elem)

print(f"  Inserted {len(toc_elems)} TOC elements")

# ══════════════════════════════════════════════════════════════
# STEP 4: Insert ภาคผนวก after บรรณานุกรม
# ══════════════════════════════════════════════════════════════
print("Step 4: Inserting appendices...")

# Find บรรณานุกรม heading and its last entry
bib_heading = None
bib_last = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'บรรณานุกรม':
        bib_heading = i
    if bib_heading and i > bib_heading and p.text.strip():
        bib_last = i

print(f"  บรรณานุกรม heading at {bib_heading}, last entry at {bib_last}")

# ── ภาคผนวก ก: Project Structure ──
project_structure = """src/
├── app/
│   ├── (protected)/
│   │   ├── assistant/page.tsx
│   │   ├── blocks/page.tsx
│   │   ├── buildings/page.tsx
│   │   ├── categories/page.tsx
│   │   ├── dashboard/page.tsx
│   │   ├── import/page.tsx
│   │   ├── movements/page.tsx
│   │   ├── parts/page.tsx
│   │   ├── parts/new/page.tsx
│   │   ├── parts/[id]/page.tsx
│   │   ├── scan/page.tsx
│   │   ├── settings/password/page.tsx
│   │   └── users/page.tsx
│   ├── api/
│   │   ├── ai/chat/route.ts
│   │   ├── ai/actions/[id]/confirm/route.ts
│   │   ├── auth/login/route.ts
│   │   ├── dashboard/route.ts
│   │   ├── export/route.ts
│   │   ├── import/route.ts
│   │   ├── line/webhook/route.ts
│   │   ├── liff/parts/route.ts
│   │   ├── mobile/ (API สำหรับแอปพลิเคชันมือถือ)
│   │   └── ... (API Routes อื่น ๆ)
│   ├── liff/
│   │   ├── add-part/page.tsx
│   │   ├── search/page.tsx
│   │   ├── scan/page.tsx
│   │   └── stock-move/page.tsx
│   └── login/page.tsx
├── components/
│   ├── layout/ (sidebar, bottom-nav, page-header)
│   └── ui/ (shadcn/ui components)
└── lib/
    ├── ai-assistant/ (orchestrator, tools, pending-actions)
    ├── line-chat/ (orchestrator, flex-messages, tools)
    ├── prisma.ts
    ├── auth.ts
    ├── excel.ts
    ├── qrcode.ts
    └── ... (ไลบรารีอื่น ๆ)"""

# ── ภาคผนวก ข: Prisma Schema ──
schema_path = os.path.join(BASE, "prisma", "schema.prisma")
schema_content = ""
if os.path.exists(schema_path):
    with open(schema_path, 'r', encoding='utf-8') as f:
        schema_content = f.read()

# ── ภาคผนวก ง: README.md ──
readme_path = os.path.join(BASE, "README.md")
readme_content = ""
if os.path.exists(readme_path):
    with open(readme_path, 'r', encoding='utf-8') as f:
        readme_content = f.read()

# ── ภาคผนวก จ: UAT ละเอียด ──
uat_detailed = [
    ["1", "เข้าสู่ระบบด้วยชื่อผู้ใช้และรหัสผ่านที่ถูกต้อง", "กรอกชื่อผู้ใช้ admin และรหัสผ่านที่ถูกต้อง กดปุ่มเข้าสู่ระบบ", "เข้าสู่ระบบสำเร็จ แสดงหน้า Dashboard", "เข้าสู่ระบบสำเร็จ แสดงหน้า Dashboard", "ผ่าน"],
    ["2", "เข้าสู่ระบบด้วยรหัสผ่านผิด", "กรอกชื่อผู้ใช้ admin และรหัสผ่านที่ผิด กดปุ่มเข้าสู่ระบบ", "แสดงข้อความแจ้งเตือนรหัสผ่านไม่ถูกต้อง", "แสดงข้อความแจ้งเตือนรหัสผ่านไม่ถูกต้อง", "ผ่าน"],
    ["3", "แสดงผลหน้า Dashboard หลังเข้าสู่ระบบ", "เข้าสู่ระบบสำเร็จ ดูข้อมูลบนหน้า Dashboard", "แสดงข้อมูลสรุปจำนวนอะไหล่ หมวดหมู่ อาคาร และอะไหล่ใกล้หมด", "แสดงข้อมูลสรุปครบถ้วน", "ผ่าน"],
    ["4", "เพิ่มข้อมูลอะไหล่ใหม่", "กรอกข้อมูลอะไหล่ใหม่ครบถ้วน กดปุ่มบันทึก", "บันทึกข้อมูลอะไหล่ใหม่สำเร็จ แสดงในรายการ", "บันทึกสำเร็จ อะไหล่ปรากฏในรายการ", "ผ่าน"],
    ["5", "เพิ่มอะไหล่ใหม่ด้วยรหัสที่ซ้ำกับที่มีอยู่", "กรอกรหัสอะไหล่ที่มีอยู่แล้วในระบบ กดปุ่มบันทึก", "แสดงข้อความแจ้งเตือนรหัสอะไหล่ซ้ำ", "แสดงข้อความแจ้งเตือนรหัสซ้ำ", "ผ่าน"],
    ["6", "แก้ไขข้อมูลอะไหล่", "เข้าหน้ารายละเอียดอะไหล่ กดแก้ไข เปลี่ยนข้อมูล กดบันทึก", "บันทึกการแก้ไขสำเร็จ ข้อมูลอัปเดต", "บันทึกสำเร็จ ข้อมูลเปลี่ยนแปลง", "ผ่าน"],
    ["7", "ค้นหาอะไหล่จากรหัสอะไหล่", "พิมพ์รหัสอะไหล่ในช่องค้นหา กดค้นหา", "แสดงรายการอะไหล่ที่ตรงกับรหัสที่ค้นหา", "แสดงผลลัพธ์ตรงกับรหัส", "ผ่าน"],
    ["8", "ค้นหาอะไหล่จากชื่ออะไหล่", "พิมพ์ชื่ออะไหล่ในช่องค้นหา กดค้นหา", "แสดงรายการอะไหล่ที่มีชื่อตรงหรือใกล้เคียง", "แสดงผลลัพธ์ใกล้เคียง", "ผ่าน"],
    ["9", "ค้นหาอะไหล่จาก Barcode", "สแกน Barcode หรือพิมพ์ค่า Barcode ในช่องค้นหา", "แสดงรายการอะไหล่ที่ตรงกับ Barcode", "แสดงผลลัพธ์ตรงกับ Barcode", "ผ่าน"],
    ["10", "แสดงสถานะอะไหล่ (มีของ/ใกล้หมด/หมดสต็อก)", "ดูสถานะอะไหล่ในรายการ", "แสดงสถานะอะไหล่ถูกต้องตามจำนวนคงเหลือและจำนวนขั้นต่ำ", "สถานะแสดงถูกต้อง", "ผ่าน"],
    ["11", "บันทึกการรับเข้าอะไหล่", "เข้าหน้า Stock Movement เลือกรับเข้า กรอกจำนวน กดบันทึก", "จำนวนอะไหล่เพิ่มขึ้น ประวัติถูกบันทึก", "จำนวนเพิ่มขึ้น ประวัติบันทึกถูกต้อง", "ผ่าน"],
    ["12", "บันทึกการเบิกออกอะไหล่", "เข้าหน้า Stock Movement เลือกเบิกออก กรอกจำนวน กดบันทึก", "จำนวนอะไหล่ลดลง ประวัติถูกบันทึก", "จำนวนลดลง ประวัติบันทึกถูกต้อง", "ผ่าน"],
    ["13", "บันทึกการเบิกออกเกินจำนวนคงเหลือ", "เบิกออกอะไหล่จำนวนเกินกว่าคงเหลือ", "ระบบป้องกันไม่ให้เบิกเกิน แสดงข้อความแจ้งเตือน", "ระบบป้องกัน แจ้งเตือนถูกต้อง", "ผ่าน"],
    ["14", "บันทึกการปรับยอดอะไหล่", "เข้าหน้า Stock Movement เลือกปรับยอด กรอกจำนวนใหม่ กดบันทึก", "จำนวนอะไหล่ถูกปรับตามจำนวนที่ระบุ ประวัติถูกบันทึก", "จำนวนถูกปรับ ประวัติบันทึกถูกต้อง", "ผ่าน"],
    ["15", "ตรวจสอบประวัติการเคลื่อนไหวของสต็อก", "เข้าหน้าประวัติ Stock Movement", "แสดงประวัติการรับเข้า เบิกออก และปรับยอดย้อนหลัง", "แสดงประวัติครบถ้วน", "ผ่าน"],
    ["16", "สแกน QR Code เพื่อเข้าถึงข้อมูลอะไหล่", "เปิดหน้าสแกน สแกน QR Code ของอะไหล่", "แสดงข้อมูลอะไหล่ที่ตรงกับ QR Code ที่สแกน", "แสดงข้อมูลอะไหล่ถูกต้อง", "ผ่าน"],
    ["17", "สแกน QR Code ที่ไม่มีในระบบ", "สแกน QR Code ที่ไม่ตรงกับอะไหล่ใดในระบบ", "แสดงข้อความแจ้งว่าไม่พบอะไหล่", "แจ้งไม่พบอะไหล่ถูกต้อง", "ผ่าน"],
    ["18", "นำเข้าข้อมูลอะไหล่จากไฟล์ Excel ที่ถูกต้อง", "อัปโหลดไฟล์ Excel ที่มีข้อมูลถูกต้อง", "นำเข้าข้อมูลสำเร็จ แสดงจำนวนรายการที่นำเข้า", "นำเข้าสำเร็จ จำนวนรายการตรง", "ผ่าน"],
    ["19", "นำเข้าข้อมูลจากไฟล์ Excel ที่มีรหัสอะไหล่ซ้ำ", "อัปโหลดไฟล์ Excel ที่มีรหัสอะไหล่ซ้ำกับที่มีอยู่", "แสดงข้อความแจ้งเตือนรายการที่ซ้ำ", "แจ้งเตือนรายการซ้ำถูกต้อง", "ผ่าน"],
    ["20", "เพิ่มหมวดหมู่อะไหล่ใหม่", "เข้าหน้าจัดการหมวดหมู่ กรอกชื่อหมวดหมู่ใหม่ กดบันทึก", "บันทึกหมวดหมู่ใหม่สำเร็จ", "บันทึกสำเร็จ", "ผ่าน"],
    ["21", "แก้ไขชื่อหมวดหมู่", "เลือกหมวดหมู่ กดแก้ไข เปลี่ยนชื่อ กดบันทึก", "บันทึกการแก้ไขสำเร็จ", "บันทึกสำเร็จ", "ผ่าน"],
    ["22", "ลบหมวดหมู่ที่ไม่มีอะไหล่อยู่", "เลือกหมวดหมู่ที่ไม่มีอะไหล่ กดลบ", "ลบหมวดหมู่สำเร็จ", "ลบสำเร็จ", "ผ่าน"],
    ["23", "เพิ่มข้อมูลอาคารใหม่", "เข้าหน้าจัดการอาคาร กรอกชื่ออาคารใหม่ กดบันทึก", "บันทึกอาคารใหม่สำเร็จ", "บันทึกสำเร็จ", "ผ่าน"],
    ["24", "จัดลำดับอาคาร", "เลือกอาคาร เปลี่ยนลำดับ กดบันทึก", "ลำดับอาคารเปลี่ยนแปลงตามที่กำหนด", "ลำดับเปลี่ยนถูกต้อง", "ผ่าน"],
    ["25", "เพิ่มผู้ใช้งานใหม่โดยผู้ดูแลระบบ", "เข้าหน้าจัดการผู้ใช้ กรอกข้อมูลผู้ใช้ใหม่ กดบันทึก", "บันทึกผู้ใช้ใหม่สำเร็จ", "บันทึกสำเร็จ", "ผ่าน"],
    ["26", "เปิด/ปิดใช้งานบัญชีผู้ใช้", "เลือกผู้ใช้ กดปุ่มเปิด/ปิดใช้งาน", "สถานะบัญชีเปลี่ยนแปลงตามที่กำหนด", "สถานะเปลี่ยนถูกต้อง", "ผ่าน"],
    ["27", "สอบถามข้อมูลอะไหล่ผ่าน AI Assistant", "เปิดหน้า AI Assistant พิมพ์คำถามเกี่ยวกับอะไหล่", "AI ตอบคำถามเกี่ยวกับอะไหล่ได้ถูกต้อง", "AI ตอบถูกต้อง", "ผ่าน"],
    ["28", "ค้นหาอะไหล่ด้วยรูปภาพผ่าน AI", "อัปโหลดรูปภาพอะไหล่ในระบบ AI", "แสดงรายการอะไหล่ที่ใกล้เคียงกับรูปภาพ", "แสดงผลลัพธ์ใกล้เคียง", "ผ่าน"],
    ["29", "สอบถามข้อมูลอะไหล่ผ่าน LINE Bot", "ส่งข้อความสอบถามอะไหล่ผ่าน LINE", "LINE Bot ตอบคำถามและแสดงข้อมูลอะไหล่ได้ถูกต้อง", "LINE Bot ตอบถูกต้อง", "ผ่าน"],
    ["30", "เข้าใช้งานระบบผ่าน LINE LIFF", "เปิด LIFF URL ในแอปพลิเคชัน LINE", "เปิดหน้าเว็บแอปภายใน LINE ได้ แสดงข้อมูลอะไหล่", "เปิดได้ แสดงข้อมูลถูกต้อง", "ผ่าน"],
    ["31", "เข้าสู่ระบบแอปพลิเคชันมือถือ", "เปิดแอปพลิเคชัน กรอกชื่อผู้ใช้และรหัสผ่าน", "เข้าสู่ระบบสำเร็จ แสดงหน้า Dashboard บนมือถือ", "เข้าสู่ระบบสำเร็จ", "ผ่าน"],
    ["32", "ค้นหาอะไหล่บนแอปพลิเคชันมือถือ", "พิมพ์คำค้นหาในแอปพลิเคชันมือถือ", "แสดงรายการอะไหล่ตามคำค้นหา", "แสดงผลลัพธ์ถูกต้อง", "ผ่าน"],
    ["33", "สแกน QR Code/Barcode บนแอปพลิเคชันมือถือ", "ใช้กล้องสแกน QR Code/Barcode ในแอป", "แสดงข้อมูลอะไหล่ที่ตรงกับรหัสที่สแกน", "แสดงข้อมูลถูกต้อง", "ผ่าน"],
    ["34", "เพิ่มอะไหล่ใหม่ผ่านแอปพลิเคชันมือถือ", "กรอกข้อมูลอะไหล่ใหม่ในแอป กดบันทึก", "บันทึกข้อมูลอะไหล่ใหม่สำเร็จจากแอปพลิเคชันมือถือ", "บันทึกสำเร็จ", "ผ่าน"],
    ["35", "ตรวจสอบสิทธิ์ผู้ใช้งานตามบทบาท (ADMIN/STAFF)", "เข้าสู่ระบบด้วยบัญชี ADMIN และ STAFF เปรียบเทียบสิทธิ์", "ผู้ดูแลระบบใช้งานฟังก์ชันจัดการได้ เจ้าหน้าที่ทั่วไปใช้งานตามสิทธิ์", "สิทธิ์แยกตามบทบาทถูกต้อง", "ผ่าน"],
]

# Build appendix elements
appendix_elems = []

# ── ภาคผนวก heading ──
appendix_elems.append(_add_page_break_elem())
appendix_elems.append(_make_para("ภาคผนวก", ref_idx=REF_CHAPTER, bold=True, alignment='center'))
appendix_elems.append(_make_para("", ref_idx=REF_BODY))

# ── ภาคผนวก ก ──
appendix_elems.append(_make_para("ภาคผนวก ก", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("โครงสร้างโปรเจกต์ (Project Structure)", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("", ref_idx=REF_BODY))

for line in project_structure.split('\n'):
    appendix_elems.append(_make_para(line, ref_idx=REF_BODY, font_size=14))

# ── ภาคผนวก ข ──
appendix_elems.append(_add_page_break_elem())
appendix_elems.append(_make_para("ภาคผนวก ข", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("โค้ดสำคัญของระบบ — Prisma Schema (prisma/schema.prisma)", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("", ref_idx=REF_BODY))

if schema_content:
    for line in schema_content.split('\n'):
        # Truncate very long lines
        if len(line) > 100:
            line = line[:97] + '...'
        appendix_elems.append(_make_para(line, ref_idx=REF_BODY, font_size=14))

# ── ภาคผนวก ง ──
appendix_elems.append(_add_page_break_elem())
appendix_elems.append(_make_para("ภาคผนวก ง", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("README.md — คู่มือการติดตั้งและใช้งานระบบ", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("", ref_idx=REF_BODY))

if readme_content:
    for line in readme_content.split('\n'):
        if len(line) > 100:
            line = line[:97] + '...'
        appendix_elems.append(_make_para(line, ref_idx=REF_BODY, font_size=14))

# ── ภาคผนวก จ ──
appendix_elems.append(_add_page_break_elem())
appendix_elems.append(_make_para("ภาคผนวก จ", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("ผลการทดสอบการยอมรับระบบ (UAT) แบบละเอียด", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("", ref_idx=REF_BODY))

# Create UAT detailed table
uat_table = _make_table_elem(
    ["ลำดับ", "รายการทดสอบ", "ขั้นตอนการทดสอบ", "ผลที่คาดหวัง", "ผลจริง", "ผลการทดสอบ"],
    uat_detailed
)
appendix_elems.append(uat_table)

# Insert all appendix elements after the last bibliography entry
if bib_last:
    ref_elem = doc.paragraphs[bib_last]._element
    for elem in reversed(appendix_elems):
        ref_elem.addnext(elem)
    print(f"  Inserted {len(appendix_elems)} appendix elements after paragraph {bib_last}")
else:
    print("  WARNING: Could not find bibliography, appending at end")
    sp = body.find(qn('w:sectPr'))
    for elem in appendix_elems:
        if sp is not None:
            sp.addprevious(elem)
        else:
            body.append(elem)

# ══════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════
print(f"\nSaving to {OUT_PATH}...")
doc.save(OUT_PATH)
print("Done!")

# ── Verify ──
doc2 = Document(OUT_PATH)
print(f"\nVerification:")
print(f"  Total paragraphs: {len(doc2.paragraphs)}")
print(f"  Total tables: {len(doc2.tables)}")

# Check TOC
toc_found = False
fig_toc_found = False
tbl_toc_found = False
for p in doc2.paragraphs:
    if p.text.strip() == 'สารบัญ':
        toc_found = True
    if p.text.strip() == 'สารบัญรูป':
        fig_toc_found = True
    if p.text.strip() == 'สารบัญตาราง':
        tbl_toc_found = True

print(f"  สารบัญ: {'OK' if toc_found else 'MISSING'}")
print(f"  สารบัญรูป: {'OK' if fig_toc_found else 'MISSING'}")
print(f"  สารบัญตาราง: {'OK' if tbl_toc_found else 'MISSING'}")

# Check appendices
app_kor_found = False
app_kor_found2 = False
app_ngoe_found = False
app_jo_found = False
for p in doc2.paragraphs:
    t = p.text.strip()
    if t == 'ภาคผนวก ก':
        app_kor_found = True
    if t == 'ภาคผนวก ข':
        app_kor_found2 = True
    if t == 'ภาคผนวก ง':
        app_ngoe_found = True
    if t == 'ภาคผนวก จ':
        app_jo_found = True

print(f"  ภาคผนวก ก: {'OK' if app_kor_found else 'MISSING'}")
print(f"  ภาคผนวก ข: {'OK' if app_kor_found2 else 'MISSING'}")
print(f"  ภาคผนวก ง: {'OK' if app_ngoe_found else 'MISSING'}")
print(f"  ภาคผนวก จ: {'OK' if app_jo_found else 'MISSING'}")

# Check figure numbering
print("\n  Figure captions after renumbering:")
for p in doc2.paragraphs:
    t = p.text.strip()
    if t.startswith('รูปที่ 3.'):
        print(f"    {t[:70]}")
