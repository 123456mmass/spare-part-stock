"""
Generate internship report v9 — improve v8 with page numbers, better code
formatting, fuller project structure, and more detailed UAT.

Changes from v8→v9:
  1. Add estimated page numbers to Table of Contents, List of Figures, List of Tables
  2. Format code appendices (Prisma schema, README) with monospace font + gray
     background so they look like VS Code code blocks
  3. Expand Project Structure appendix with root-level files and complete src tree
  4. Expand UAT appendix with step-by-step detailed test procedures
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import os, copy, re

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMG_DIR = os.path.join(BASE, "docs", "images")
INPUT = os.path.join(BASE, "docs", "report-v7-final.docx")
OUT_PATH = os.path.join(BASE, "docs", "report-v9-final-NEW.docx")

FONT = 'TH SarabunPSK'

doc = Document(INPUT)
body = doc.element.body

# ── Reference paragraph indices ──
# Find useful reference paragraphs for cloning
REF_CHAPTER = None   # e.g. "บทที่ 1 บทนำ"
REF_SECTION = None   # e.g. "1.1 ความเป็นมา..."
REF_BODY = None      # normal body text
REF_CAPTION = None   # italic caption like "รูปที่ ..."

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if REF_CHAPTER is None and t.startswith('บทที่'):
        REF_CHAPTER = i
    if REF_SECTION is None and re.match(r'^\d+\.\d+\s', t):
        REF_SECTION = i
    if REF_BODY is None and len(t) > 40 and not t.startswith('บทที่') and not t.startswith('รูปที่'):
        REF_BODY = i
    if REF_CAPTION is None and t.startswith('รูปที่'):
        REF_CAPTION = i
    if all([REF_CHAPTER, REF_SECTION, REF_BODY, REF_CAPTION]):
        break

print(f"Reference paragraphs: chapter={REF_CHAPTER}, section={REF_SECTION}, body={REF_BODY}, caption={REF_CAPTION}")

# ══════════════════════════════════════════════════════════════
# Helper functions
# ══════════════════════════════════════════════════════════════
def _escape_xml(text):
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')

def _make_para(text, ref_idx=None, bold=False, alignment=None, italic=False, font_size=None):
    if ref_idx is None:
        ref_idx = REF_BODY or 64
    ref_elem = doc.paragraphs[ref_idx]._element
    new_p = copy.deepcopy(ref_elem)
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    ref_runs = ref_elem.findall(qn('w:r'))
    new_r = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
    if ref_runs:
        ref_rPr = ref_runs[0].find(qn('w:rPr'))
        if ref_rPr is not None:
            new_rPr = copy.deepcopy(ref_rPr)
            if bold:
                for tag in ['w:b', 'w:bCs']:
                    e = new_rPr.find(qn(tag))
                    if e is None:
                        e = parse_xml(f'<{tag} {nsdecls("w")}/>')
                        new_rPr.append(e)
                    else:
                        e.set(qn('w:val'), '1')
            else:
                for tag in ['w:b', 'w:bCs']:
                    e = new_rPr.find(qn(tag))
                    if e is not None:
                        e.set(qn('w:val'), '0')
            if italic:
                for tag in ['w:i', 'w:iCs']:
                    e = new_rPr.find(qn(tag))
                    if e is None:
                        e = parse_xml(f'<{tag} {nsdecls("w")}/>')
                        new_rPr.append(e)
                    else:
                        e.set(qn('w:val'), '1')
            if font_size:
                for tag in ['w:sz', 'w:szCs']:
                    e = new_rPr.find(qn(tag))
                    if e is not None:
                        e.set(qn('w:val'), str(font_size * 2))
            new_r.append(new_rPr)
    t = parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve">{_escape_xml(text)}</w:t>')
    new_r.append(t)
    new_p.append(new_r)
    if alignment:
        pPr = new_p.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            new_p.insert(0, pPr)
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="{alignment}"/>'))
        else:
            jc.set(qn('w:val'), alignment)
    return new_p

def _add_page_break_elem():
    return parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:r><w:br w:type="page"/></w:r>'
        f'</w:p>'
    )

def _make_code_para(lines, font_name='Consolas', font_size=10, shading='F2F2F2'):
    """Create a code-block style paragraph (monospaced, small font, light gray bg).
    Multiple lines are kept inside ONE paragraph separated by line breaks,
    exactly like a VS Code code block.
    """
    if isinstance(lines, str):
        lines = lines.split('\n')
    # Use body paragraph as structural base
    ref_elem = doc.paragraphs[REF_BODY]._element
    new_p = copy.deepcopy(ref_elem)
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)

    new_r = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
    new_rPr = parse_xml(
        f'<w:rPr {nsdecls("w")}>'
        f'  <w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}" w:eastAsia="{font_name}"/>'
        f'  <w:sz w:val="{font_size * 2}"/>'
        f'  <w:szCs w:val="{font_size * 2}"/>'
        f'</w:rPr>'
    )
    new_r.append(new_rPr)

    for i, line in enumerate(lines):
        if i > 0:
            new_r.append(parse_xml(f'<w:br {nsdecls("w")}/>'))
        t = parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve">{_escape_xml(line)}</w:t>')
        new_r.append(t)

    new_p.append(new_r)

    # Add light gray background shading to paragraph
    pPr = new_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        new_p.insert(0, pPr)
    shd = pPr.find(qn('w:shd'))
    if shd is None:
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{shading}" w:color="auto"/>')
        pPr.append(shd)
    else:
        shd.set(qn('w:fill'), shading)

    # Remove spacing between lines for compact code look
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>')
        pPr.append(spacing)
    else:
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')

    # Add a little left indent for code block visual
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="284"/>')
        pPr.append(ind)
    else:
        ind.set(qn('w:left'), '284')

    return new_p

def _add_image_after(ref_elem, filename, width=Inches(5.2)):
    path = os.path.join(IMG_DIR, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if os.path.exists(path):
        run.add_picture(path, width=width)
    else:
        run.text = f"[รูปภาพ: {filename}]"
    p_elem = p._element
    body.remove(p_elem)
    ref_elem.addnext(p_elem)
    return p_elem

def _make_table_elem(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = t._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = FONT
                r.font.size = Pt(18)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            val_str = str(val)
            # Render multi-line values as multiple paragraphs inside the cell
            if '\n' in val_str:
                lines = val_str.split('\n')
                cell.text = lines[0]
                for extra in lines[1:]:
                    p = cell.add_paragraph()
                    p.text = extra
            else:
                cell.text = val_str
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT
                    r.font.size = Pt(18)
    tbl_elem = t._tbl
    tbl_elem.getparent().remove(tbl_elem)
    return tbl_elem

def _dots(label, page=""):
    """Generate TOC line with dot leaders and right-aligned page number."""
    if page:
        total_width = 72
        sep = f" {page}"
        # Account for multi-byte Thai characters by using display width estimate
        label_display = len(label)
        dot_count = max(3, total_width - label_display - len(sep))
        return f"{label}{'.' * dot_count}{sep}"
    return label

def _estimate_page_number(para_idx, base_para_idx=0, blank_factor=0.7, lines_per_page=48):
    """Estimate page number from paragraph index.
    Counts non-empty paragraphs as lines, with images/tables counted as extra.
    """
    line_count = 0
    for i in range(base_para_idx, min(para_idx, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if not text:
            continue
        # Image captions, tables, and headings consume more page space
        if text.startswith('รูปที่') or text.startswith('ตารางที่'):
            line_count += 2
        elif text.startswith('บทที่') or re.match(r'^\d+\.\d+\s', text):
            line_count += 1
        else:
            # Body lines wrap within page width; treat paragraph as ~0.8 of a line
            line_count += max(1, len(text) // 90)
    page = max(1, 1 + int(line_count // lines_per_page))
    return str(page)

# ══════════════════════════════════════════════════════════════
# STEP 1: Fix screenshot figure numbers 3.1–3.19 → 3.5–3.23
# ══════════════════════════════════════════════════════════════
print("Step 1: Fixing screenshot figure numbers...")

DIAGRAM_KEYWORDS = ['ER Diagram', 'System Architecture', 'แผนผังลำดับการทำงาน']
fix_count = 0

# Build a map of screenshot captions that need renumbering
# Diagrams keep 3.1-3.4, screenshots shift by +4
# We need to process ALL runs of a paragraph together, because the
# "รูปที่ 3.16" might be split across runs like "รูปที่ 3.1" + "6"

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t.startswith('รูปที่'):
        continue
    # Skip diagrams — they keep their numbers
    if any(k in t for k in DIAGRAM_KEYWORDS):
        continue
    # This is a screenshot caption — fix the number
    # Extract the full number from the combined text
    m = re.match(r'รูปที่\s*3\.(\d+)', t)
    if not m:
        continue
    old_num = int(m.group(1))
    new_num = old_num + 4
    old_str = f'3.{old_num}'
    new_str = f'3.{new_num}'

    # Find which run(s) contain the number and replace
    # Simple approach: just replace in any run that contains part of the number
    for run in p.runs:
        if old_str in run.text:
            run.text = run.text.replace(old_str, new_str)
            fix_count += 1
            print(f"  [{i}] {t[:50]} -> รูปที่ {new_str}")
            break
    else:
        # Number might be split across runs — try combining first two runs
        if len(p.runs) >= 2:
            combined = p.runs[0].text + p.runs[1].text
            if old_str in combined:
                # Put the full replacement in the first run, clear the second's number part
                p.runs[0].text = p.runs[0].text.replace(old_str, new_str)
                if p.runs[1].text.startswith(str(old_num)[-1:]):
                    p.runs[1].text = p.runs[1].text[len(str(old_num))-len(str(old_num).rstrip(str(old_num)[-1])):]
                fix_count += 1
                print(f"  [{i}] (split) {t[:50]} -> รูปที่ {new_str}")

print(f"  Fixed {fix_count} screenshot captions")

# ══════════════════════════════════════════════════════════════
# STEP 2: Collect all headings, figures, tables for TOC
# ══════════════════════════════════════════════════════════════
print("Step 2: Collecting TOC entries...")

toc_entries = []  # (level, text, page_placeholder)
fig_entries = []  # (fig_num, caption)
tbl_entries = []  # (tbl_num, caption)

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue

    # Chapters
    if t.startswith('บทที่'):
        toc_entries.append((0, t, ""))
    # Sections like 1.1, 2.3, 3.6
    elif re.match(r'^\d+\.\d+\s', t) and len(t) < 80:
        # Sub-sections like 1.3.1
        if re.match(r'^\d+\.\d+\.\d+\s', t):
            toc_entries.append((2, t, ""))
        else:
            toc_entries.append((1, t, ""))
    # Figure captions
    elif t.startswith('รูปที่'):
        fig_entries.append(t)
    # Standalone section headings like สารบัญ, บรรณานุกรม, ภาคผนวก
    elif t in ['สารบัญ', 'สารบัญรูป', 'สารบัญตาราง', 'บรรณานุกรม', 'ภาคผนวก']:
        toc_entries.append((0, t, ""))

# Table captions (manual — based on our knowledge of the document)
table_captions = [
    "ตารางที่ 1.1 ตัวชี้วัดประสิทธิภาพ (KPI) ก่อนและหลังการพัฒนาระบบ",
    "ตารางที่ 1.2 นิยามศัพท์เฉพาะ",
    "ตารางที่ 3.1 โครงสร้างตารางฐานข้อมูลของระบบ",
    "ตารางที่ 3.2 ผลการทดสอบการยอมรับระบบ (UAT)",
]

print(f"  Found {len(toc_entries)} TOC entries, {len(fig_entries)} figures, {len(table_captions)} tables")

# ══════════════════════════════════════════════════════════════
# STEP 3: Insert สารบัญ + สารบัญรูป + สารบัญตาราง
#         after หนังสือรับรอง (before บทที่ 1)
# ══════════════════════════════════════════════════════════════
print("Step 3: Inserting TOC pages...")

# Find the paragraph just before บทที่ 1
ch1_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('บทที่ 1'):
        ch1_idx = i
        break

if ch1_idx is None:
    print("  ERROR: Could not find บทที่ 1")
    sys.exit(1)

# Insert before บทที่ 1 by inserting after the paragraph just before it
# Find the last non-empty paragraph before บทที่ 1
pre_ch1 = ch1_idx - 1
while pre_ch1 > 0 and not doc.paragraphs[pre_ch1].text.strip():
    pre_ch1 -= 1

print(f"  บทที่ 1 at paragraph {ch1_idx}, inserting TOC after paragraph {pre_ch1}")

# Build all TOC elements, then insert in reverse order before บทที่ 1
toc_elems = []

# --- สารบัญ ---
toc_elems.append(_add_page_break_elem())
toc_title = _make_para("สารบัญ", ref_idx=REF_CHAPTER, bold=True, alignment='center')
toc_elems.append(toc_title)
toc_elems.append(_make_para("", ref_idx=REF_BODY))  # blank line

# Add TOC entries
for level, text, _ in toc_entries:
    if level == 0:
        # Find page of this chapter by searching paragraph index
        target_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith(text.split('—')[0].strip()):
                target_idx = i
                break
        page = _estimate_page_number(target_idx, base_para_idx=ch1_idx) if target_idx else ""
        # Special cases: สารบัญ etc. are before chapter 1
        if text in ['สารบัญ', 'สารบัญรูป', 'สารบัญตาราง']:
            page = "6" if text == 'สารบัญ' else "7"
        elif text == 'บรรณานุกรม':
            page = "40"
        elif text == 'ภาคผนวก':
            page = "41"
        line = _dots(text, page)
        entry = _make_para(line, ref_idx=REF_SECTION, bold=True)
    elif level == 1:
        target_idx = None
        search_prefix = text.split()[0]
        for i, p in enumerate(doc.paragraphs):
            pt = p.text.strip()
            if pt.startswith(text):
                target_idx = i
                break
        page = _estimate_page_number(target_idx, base_para_idx=ch1_idx) if target_idx else ""
        line = _dots(f"     {text}", page)
        entry = _make_para(line, ref_idx=REF_BODY)
    else:
        target_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith(text):
                target_idx = i
                break
        page = _estimate_page_number(target_idx, base_para_idx=ch1_idx) if target_idx else ""
        line = _dots(f"          {text}", page)
        entry = _make_para(line, ref_idx=REF_BODY, font_size=16)
    toc_elems.append(entry)

# --- สารบัญรูป ---
toc_elems.append(_add_page_break_elem())
fig_title = _make_para("สารบัญรูป", ref_idx=REF_CHAPTER, bold=True, alignment='center')
toc_elems.append(fig_title)
toc_elems.append(_make_para("", ref_idx=REF_BODY))

for fig_text in fig_entries:
    # Find figure caption paragraph
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(fig_text):
            target_idx = i
            break
    page = _estimate_page_number(target_idx, base_para_idx=ch1_idx) if target_idx else ""
    line = _dots(fig_text, page)
    entry = _make_para(line, ref_idx=REF_BODY)
    toc_elems.append(entry)

# --- สารบัญตาราง ---
toc_elems.append(_add_page_break_elem())
tbl_title = _make_para("สารบัญตาราง", ref_idx=REF_CHAPTER, bold=True, alignment='center')
toc_elems.append(tbl_title)
toc_elems.append(_make_para("", ref_idx=REF_BODY))

# Map table captions to their approximate page positions
table_pages = {
    "ตารางที่ 1.1": "4",
    "ตารางที่ 1.2": "11",
    "ตารางที่ 3.1": "27",
    "ตารางที่ 3.2": "39",
}

for tbl_cap in table_captions:
    prefix = ""
    for k in table_pages:
        if tbl_cap.startswith(k):
            prefix = k
            break
    page = table_pages.get(prefix, "")
    line = _dots(tbl_cap, page)
    entry = _make_para(line, ref_idx=REF_BODY)
    toc_elems.append(entry)

# Insert all elements before บทที่ 1 (by inserting after pre_ch1 in reverse)
ref_elem = doc.paragraphs[pre_ch1]._element
for elem in reversed(toc_elems):
    ref_elem.addnext(elem)

print(f"  Inserted {len(toc_elems)} TOC elements")

# ══════════════════════════════════════════════════════════════
# STEP 4: Insert ภาคผนวก after บรรณานุกรม
# ══════════════════════════════════════════════════════════════
print("Step 4: Inserting appendices...")

# Find บรรณานุกรม heading and its last entry
bib_heading = None
bib_last = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'บรรณานุกรม':
        bib_heading = i
    if bib_heading and i > bib_heading and p.text.strip():
        bib_last = i

print(f"  บรรณานุกรม heading at {bib_heading}, last entry at {bib_last}")

# ── ภาคผนวก ก: Project Structure ──
project_structure = """.
├── docs/                          # เอกสารรายงานและภาพประกอบ
├── prisma/
│   ├── migrations/                # Migration ฐานข้อมูล SQLite
│   └── schema.prisma              # Prisma Schema
├── public/                        # ไฟล์ static เช่น QR code, รูปภาพ
├── scripts/                       # สคริปต์สร้างรายงาน/นำเข้าข้อมูล
├── src/
│   ├── app/
│   │   ├── (protected)/           # หน้าที่ต้องล็อกอิน
│   │   │   ├── assistant/page.tsx     # AI Assistant
│   │   │   ├── blocks/page.tsx        # จัดการบล็อก
│   │   │   ├── buildings/page.tsx     # จัดการอาคาร
│   │   │   ├── categories/page.tsx    # จัดการหมวดหมู่
│   │   │   ├── dashboard/page.tsx     # สรุป Dashboard
│   │   │   ├── import/page.tsx        # นำเข้า Excel
│   │   │   ├── movements/page.tsx     # ประวัติ Stock Movement
│   │   │   ├── parts/
│   │   │   │   ├── new/page.tsx       # เพิ่มอะไหล่ใหม่
│   │   │   │   ├── [id]/page.tsx      # รายละเอียดอะไหล่
│   │   │   │   └── page.tsx           # รายการอะไหล่
│   │   │   ├── scan/page.tsx          # สแกน QR/Barcode
│   │   │   ├── settings/
│   │   │   │   ├── ai/page.tsx        # ตั้งค่า AI
│   │   │   │   └── password/page.tsx  # เปลี่ยนรหัสผ่าน
│   │   │   └── users/page.tsx         # จัดการผู้ใช้
│   │   ├── api/                   # Next.js API Routes
│   │   │   ├── account/           # ข้อมูลบัญชีผู้ใช้
│   │   │   ├── admin/             # APIs สำหรับผู้ดูแลระบบ
│   │   │   ├── ai/                # AI: chat, actions
│   │   │   ├── auth/              # Authentication
│   │   │   ├── blocks/
│   │   │   ├── buildings/
│   │   │   ├── categories/
│   │   │   ├── cron/
│   │   │   ├── dashboard/
│   │   │   ├── export/
│   │   │   ├── import/
│   │   │   ├── liff/              # LINE LIFF APIs
│   │   │   ├── line/              # LINE Webhook
│   │   │   ├── mobile/            # REST APIs สำหรับ App มือถือ
│   │   │   ├── movements/
│   │   │   └── parts/
│   │   ├── liff/                  # LINE LIFF mini app
│   │   │   ├── add-part/page.tsx
│   │   │   ├── search/page.tsx
│   │   │   └── stock-move/page.tsx
│   │   ├── login/page.tsx         # หน้าเข้าสู่ระบบ
│   │   ├── page.tsx               # หน้าแรก (redirect ไป login)
│   │   ├── favicon.ico            # Favicon
│   │   ├── globals.css            # Global styles (Tailwind v4)
│   │   └── layout.tsx             # Root layout
│   ├── components/
│   │   ├── layout/                # Sidebar, Header, BottomNav
│   │   └── ui/                    # shadcn/ui (Button, Table, Dialog, ...)
│   ├── lib/
│   │   ├── ai-assistant/          # AI Assistant engine
│   │   ├── line-chat/             # LINE Bot orchestration
│   │   ├── ai-client.ts           # LLM client config
│   │   ├── auth.ts                # Auth helpers
│   │   ├── embeddings.ts          # Vector / text embeddings
│   │   ├── excel.ts               # Excel read/write
│   │   ├── part-lookup.ts         # Part search service
│   │   ├── prisma.ts              # Prisma Client singleton
│   │   ├── session.ts             # JWT/session
│   │   ├── stock.ts               # Stock movement logic
│   │   └── utils.ts               # Utility functions
│   └── proxy.ts                   # Proxy utilities
├── .env.example                   # ตัวอย่าง environment variables
├── next.config.ts                 # ค่า config Next.js
├── package.json                   # Dependencies + scripts
├── postcss.config.mjs             # PostCSS config
├── README.md                      # คู่มือการใช้งาน
└── tsconfig.json                  # TypeScript config"""

# ── ภาคผนวก ข: Prisma Schema ──
schema_path = os.path.join(BASE, "prisma", "schema.prisma")
schema_content = ""
if os.path.exists(schema_path):
    with open(schema_path, 'r', encoding='utf-8') as f:
        schema_content = f.read()

# ── ภาคผนวก ง: README.md ──
readme_path = os.path.join(BASE, "README.md")
readme_content = ""
if os.path.exists(readme_path):
    with open(readme_path, 'r', encoding='utf-8') as f:
        readme_content = f.read()

# ── ภาคผนวก จ: UAT ละเอียด ──
# More detailed test steps with multi-line procedures
uat_detailed = [
    ["1", "เข้าสู่ระบบด้วยชื่อผู้ใช้และรหัสผ่านที่ถูกต้อง",
     "1) เปิดหน้า login\n2) กรอก username: admin\n3) กรอกรหัสผ่านที่ถูกต้อง\n4) กดปุ่ม เข้าสู่ระบบ",
     "เข้าสู่ระบบสำเร็จและ redirect ไปหน้า Dashboard", "เข้าสู่ระบบสำเร็จ", "ผ่าน"],
    ["2", "เข้าสู่ระบบด้วยรหัสผ่านผิด",
     "1) เปิดหน้า login\n2) กรอก username: admin\n3) กรอกรหัสผ่านผิด\n4) กดปุ่ม เข้าสู่ระบบ",
     "แสดงข้อความแจ้งเตือนรหัสผ่านไม่ถูกต้องและไม่ให้เข้าระบบ", "แสดงข้อความแจ้งเตือน", "ผ่าน"],
    ["3", "แสดงผลหน้า Dashboard หลังเข้าสู่ระบบ",
     "1) เข้าสู่ระบบสำเร็จ\n2) สังเกตหน้า Dashboard\n3) ตรวจสอบ card สรุปจำนวนอะไหล่, หมวดหมู่, อาคาร, อะไหล่ใกล้หมด",
     "แสดงข้อมูลสรุปครบถ้วนและถูกต้องตามฐานข้อมูล", "แสดงข้อมูลครบถ้วน", "ผ่าน"],
    ["4", "เพิ่มข้อมูลอะไหล่ใหม่",
     "1) ไปที่เมนู อะไหล่\n2) กด เพิ่มอะไหล่\n3) กรอก Part Number, Part Name, Category, Building, Quantity, Minimum Quantity\n4) กด บันทึก",
     "บันทึกสำเร็จ อะไหล่ปรากฏในรายการและมีสถานะหลังบันทึก", "บันทึกสำเร็จ", "ผ่าน"],
    ["5", "เพิ่มอะไหล่ใหม่ด้วยรหัสที่ซ้ำกับที่มีอยู่",
     "1) ไปที่เมนู อะไหล่\n2) กด เพิ่มอะไหล่\n3) กรอก Part Number ซ้ำกับที่มีอยู่\n4) กด บันทึก",
     "ระบบแสดงข้อความแจ้งเตือน รหัสอะไหล่ซ้ำ และไม่บันทึกข้อมูล", "แจ้งเตือนซ้ำ", "ผ่าน"],
    ["6", "แก้ไขข้อมูลอะไหล่",
     "1) คลิกที่อะไหล่รายการหนึ่ง\n2) กดปุ่ม แก้ไข\n3) เปลี่ยนชื่อหรือจำนวนขั้นต่ำ\n4) กด บันทึก",
     "บันทึกการแก้ไขสำเร็จ ข้อมูลอัปเดตในหน้ารายละเอียดและรายการ", "แก้ไขสำเร็จ", "ผ่าน"],
    ["7", "ค้นหาอะไหล่จากรหัสอะไหล่",
     "1) ไปที่เมนู อะไหล่\n2) พิมพ์รหัสอะไหล่ (เช่น P-1001) ในช่องค้นหา\n3) กด Enter หรือปุ่มค้นหา",
     "แสดงเฉพาะรายการอะไหล่ที่ตรงกับรหัสที่ค้นหา", "แสดงผลลัพธ์ตรง", "ผ่าน"],
    ["8", "ค้นหาอะไหล่จากชื่ออะไหล่",
     "1) ไปที่เมนู อะไหล่\n2) พิมพ์ชื่ออะไหล่บางส่วน (เช่น มอเตอร์)\n3) กด Enter หรือปุ่มค้นหา",
     "แสดงรายการอะไหล่ที่มีชื่อตรงหรือใกล้เคียงกับคำค้นหา", "แสดงผลลัพธ์ใกล้เคียง", "ผ่าน"],
    ["9", "ค้นหาอะไหล่จาก Barcode",
     "1) เปิดหน้า สแกน\n2) อ่าน Barcode จากอะไหล่ หรือพิมพ์ barcode ในช่องค้นหา\n3) กดค้นหา",
     "แสดงรายละเอียดอะไหล่ที่ตรงกับ Barcode นั้น", "แสดงผลลัพธ์ตรง", "ผ่าน"],
    ["10", "แสดงสถานะอะไหล่ (มีของ/ใกล้หมด/หมดสต็อก)",
     "1) ไปที่เมนู อะไหล่\n2) สังเกต badge สถานะของแต่ละรายการ\n3) เปรียบเทียบกับจำนวนคงเหลือและขั้นต่ำ",
     "สถานะอะไหล่แสดงถูกต้อง: มีของ พอ, ใกล้หมด น้อยกว่าขั้นต่ำ, หมดสต็อก 0", "สถานะถูกต้อง", "ผ่าน"],
    ["11", "บันทึกการรับเข้าอะไหล่",
     "1) ไปที่เมนู เคลื่อนไหวสต็อก\n2) เลือกประเภท รับเข้า\n3) เลือกอะไหล่ กรอกจำนวนและหมายเหตุ\n4) กด บันทึก",
     "จำนวนอะไหล่คงเหลือเพิ่มขึ้นตามจำนวนที่รับเข้า และมีประวัติการทำรายการ", "บันทึกสำเร็จ", "ผ่าน"],
    ["12", "บันทึกการเบิกออกอะไหล่",
     "1) ไปที่เมนู เคลื่อนไหวสต็อก\n2) เลือกประเภท เบิกออก\n3) เลือกอะไหล่ กรอกจำนวนและหมายเหตุ\n4) กด บันทึก",
     "จำนวนอะไหล่คงเหลือลดลงตามจำนวนที่เบิกออก และมีประวัติการทำรายการ", "บันทึกสำเร็จ", "ผ่าน"],
    ["13", "บันทึกการเบิกออกเกินจำนวนคงเหลือ",
     "1) ไปที่เมนู เคลื่อนไหวสต็อก\n2) เลือกประเภท เบิกออก\n3) เลือกอะไหล่ที่มีจำนวนคงเหลือน้อย\n4) กรอกจำนวนเบิกมากกว่าคงเหลือ กด บันทึก",
     "ระบบปฏิเสธการทำรายการ แสดงข้อความเตือนว่าเบิกเกินจำนวนคงเหลือ", "ป้องกันได้", "ผ่าน"],
    ["14", "บันทึกการปรับยอดอะไหล่",
     "1) ไปที่เมนู เคลื่อนไหวสต็อก\n2) เลือกประเภท ปรับยอด\n3) เลือกอะไหล่ กรอกจำนวนใหม่และเหตุผล\n4) กด บันทึก",
     "จำนวนคงเหลือเปลี่ยนเป็นจำนวนที่ปรับ และมีประวัติการปรับยอด", "บันทึกสำเร็จ", "ผ่าน"],
    ["15", "ตรวจสอบประวัติการเคลื่อนไหวของสต็อก",
     "1) ไปที่เมนู ประวัติ หรือ Movements\n2) ตรวจสอบรายการในตาราง\n3) ลองกรองตามประเภทหรือช่วงวันที่",
     "แสดงประวัติรับเข้า/เบิกออก/ปรับยอด พร้อมวันเวลาและผู้ทำรายการ เรียงจากใหม่ไปเก่า", "แสดงครบถ้วน", "ผ่าน"],
    ["16", "สแกน QR Code เพื่อเข้าถึงข้อมูลอะไหล่",
     "1) เปิดหน้า สแกน\n2) อนุญาตให้ใช้กล้อง\n3) สแกน QR Code บนอะไหล่",
     "แสดงหน้ารายละเอียดอะไหล่ที่ตรงกับ QR Code ที่สแกน", "สแกนสำเร็จ", "ผ่าน"],
    ["17", "สแกน QR Code ที่ไม่มีในระบบ",
     "1) เปิดหน้า สแกน\n2) สแกน QR Code ที่ไม่ได้ลงทะเบียน\n3) รอผลลัพธ์",
     "แสดงข้อความ ไม่พบอะไหล่ และไม่เกิด error", "แจ้งไม่พบอะไหล่", "ผ่าน"],
    ["18", "นำเข้าข้อมูลอะไหล่จากไฟล์ Excel ที่ถูกต้อง",
     "1) ไปที่เมนู นำเข้า\n2) เลือกไฟล์ Excel ที่มีคอลัมน์ Part Number, Part Name, Quantity\n3) กด นำเข้า\n4) ตรวจสอบสรุปผล",
     "นำเข้าสำเร็จ แสดงจำนวนรายการที่นำเข้าและข้อมูลปรากฏในระบบ", "นำเข้าสำเร็จ", "ผ่าน"],
    ["19", "นำเข้าข้อมูลจากไฟล์ Excel ที่มีรหัสอะไหล่ซ้ำ",
     "1) เตรียมไฟล์ Excel ที่มี Part Number ซ้ำกับที่มีอยู่\n2) ไปที่เมนู นำเข้า\n3) อัปโหลดไฟล์\n4) ตรวจสอบผล",
     "ระบบแจ้งเตือนรายการซ้ำและไม่บันทึกข้อมูลซ้ำซ้อน", "แจ้งเตือนซ้ำ", "ผ่าน"],
    ["20", "เพิ่มหมวดหมู่อะไหล่ใหม่",
     "1) ไปที่เมนู หมวดหมู่\n2) กด เพิ่มหมวดหมู่\n3) กรอกชื่อหมวดหมู่\n4) กด บันทึก",
     "บันทึกหมวดหมู่ใหม่สำเร็จ แสดงในรายการหมวดหมู่", "บันทึกสำเร็จ", "ผ่าน"],
    ["21", "แก้ไขชื่อหมวดหมู่",
     "1) ไปที่เมนู หมวดหมู่\n2) เลือกหมวดหมู่ที่ต้องการ\n3) กด แก้ไข\n4) เปลี่ยนชื่อ กด บันทึก",
     "ชื่อหมวดหมู่อัปเดตสำเร็จและแสดงผลใหม่", "แก้ไขสำเร็จ", "ผ่าน"],
    ["22", "ลบหมวดหมู่ที่ไม่มีอะไหล่อยู่",
     "1) ไปที่เมนู หมวดหมู่\n2) เลือกหมวดหมู่ที่ไม่มีอะไหล่\n3) กด ลบ\n4) ยืนยันการลบ",
     "ลบสำเร็จ หมวดหมู่หายไปจากรายการ", "ลบสำเร็จ", "ผ่าน"],
    ["23", "เพิ่มข้อมูลอาคารใหม่",
     "1) ไปที่เมนู อาคาร\n2) กด เพิ่มอาคาร\n3) กรอกชื่ออาคาร\n4) กด บันทึก",
     "บันทึกอาคารใหม่สำเร็จ แสดงในรายการอาคาร", "บันทึกสำเร็จ", "ผ่าน"],
    ["24", "จัดลำดับอาคาร",
     "1) ไปที่เมนู อาคาร\n2) ใช้ปุ่มลูกศรขึ้น/ลง หรือ drag เปลี่ยนลำดับ\n3) กด บันทึกลำดับ",
     "ลำดับอาคารเปลี่ยนแปลงและบันทึกลงฐานข้อมูล", "บันทึกลำดับสำเร็จ", "ผ่าน"],
    ["25", "เพิ่มผู้ใช้งานใหม่โดยผู้ดูแลระบบ",
     "1) เข้าสู่ระบบด้วยบัญชี ADMIN\n2) ไปที่เมนู ผู้ใช้\n3) กด เพิ่มผู้ใช้\n4) กรอก username, ชื่อ, รหัสผ่าน, บทบาท\n5) กด บันทึก",
     "บันทึกผู้ใช้ใหม่สำเร็จ ผู้ใช้สามารถล็อกอินได้", "บันทึกสำเร็จ", "ผ่าน"],
    ["26", "เปิด/ปิดใช้งานบัญชีผู้ใช้",
     "1) เข้าสู่ระบบด้วยบัญชี ADMIN\n2) ไปที่เมนู ผู้ใช้\n3) เลือกผู้ใช้\n4) กดเปิด/ปิดการใช้งาน (toggle isActive)",
     "สถานะบัญชีเปลี่ยนแปลงตามที่กำหนด (เปิด/ปิดการใช้งาน)", "เปลี่ยนสถานะสำเร็จ", "ผ่าน"],
    ["27", "สอบถามข้อมูลอะไหล่ผ่าน AI Assistant",
     "1) ไปที่เมนู AI Assistant\n2) พิมพ์คำถาม เช่น อะไหล่หมดสต็อกมีอะไรบ้าง\n3) กดส่ง\n4) รอ AI ตอบ",
     "AI ตอบคำถามด้วยข้อมูลที่ถูกต้องจากฐานข้อมูล", "AI ตอบถูกต้อง", "ผ่าน"],
    ["28", "ค้นหาอะไหล่ด้วยรูปภาพผ่าน AI",
     "1) ไปที่เมนู AI Assistant\n2) อัปโหลดรูปภาพอะไหล่\n3) รอระบบวิเคราะห์\n4) ตรวจสอบผลลัพธ์",
     "แสดงรายการอะไหล่ที่มีลักษณะใกล้เคียงกับรูปภาพ", "ค้นหาด้วยรูปได้", "ผ่าน"],
    ["29", "สอบถามข้อมูลอะไหล่ผ่าน LINE Bot",
     "1) เพิ่ม LINE Bot เป็นเพื่อน\n2) ส่งข้อความถามเกี่ยวกับอะไหล่\n3) รอการตอบกลับ",
     "LINE Bot ตอบกลับด้วยข้อมูลอะไหล่ที่ถูกต้อง", "LINE Bot ตอบถูกต้อง", "ผ่าน"],
    ["30", "เข้าใช้งานระบบผ่าน LINE LIFF",
     "1) เปิดลิงก์ LIFF ในแอป LINE\n2) ล็อกอิน (ถ้ามี)\n3) เข้าหน้าเช็คสต็อกหรือเพิ่มอะไหล่",
     "LIFF โหลดหน้าเว็บภายใน LINE ได้สำเร็จ ข้อมูลอะไหล่แสดงถูกต้อง", "LIFF ทำงานได้", "ผ่าน"],
    ["31", "เข้าสู่ระบบแอปพลิเคชันมือถือ",
     "1) เปิดแอป Flutter บนมือถือ\n2) กรอก username และรหัสผ่าน\n3) กด เข้าสู่ระบบ",
     "เข้าสู่ระบบสำเร็จ แสดงหน้า Dashboard บนมือถือ", "เข้าระบบสำเร็จ", "ผ่าน"],
    ["32", "ค้นหาอะไหล่บนแอปพลิเคชันมือถือ",
     "1) เปิดแอปบนมือถือ\n2) แตะช่องค้นหา\n3) พิมพ์คำค้นหา (ชื่อ/รหัส)\n4) แตะค้นหา",
     "แสดงรายการอะไหล่ตามคำค้นหา สามารถเลือกดูรายละเอียดได้", "ค้นหาได้", "ผ่าน"],
    ["33", "สแกน QR Code/Barcode บนแอปพลิเคชันมือถือ",
     "1) เปิดแอปบนมือถือ\n2) เข้าหน้าสแกน\n3) อนุญาตกล้อง\n4) สแกน QR/Barcode",
     "แสดงข้อมูลอะไหล่ที่ตรงกับรหัสที่สแกน", "สแกนสำเร็จ", "ผ่าน"],
    ["34", "เพิ่มอะไหล่ใหม่ผ่านแอปพลิเคชันมือถือ",
     "1) เปิดแอปบนมือถือ\n2) เข้าหน้าเพิ่มอะไหล่\n3) กรอกข้อมูลและถ่ายรูป (ถ้ามี)\n4) กด บันทึก",
     "บันทึกข้อมูลอะไหล่ใหม่สำเร็จจากแอปพลิเคชันมือถือ", "บันทึกสำเร็จ", "ผ่าน"],
    ["35", "ตรวจสอบสิทธิ์ผู้ใช้งานตามบทบาท (ADMIN/STAFF)",
     "1) เข้าสู่ระบบด้วย ADMIN: ตรวจสอบเมนูจัดการผู้ใช้/ระบบ\n2) ออกจากระบบ\n3) เข้าสู่ระบบด้วย STAFF: ตรวจสอบเมนูที่แสดง",
     "ADMIN เห็นเมนูจัดการระบบและผู้ใช้ STAFF เห็นเฉพาะฟังก์ชันตามสิทธิ์", "สิทธิ์ถูกต้อง", "ผ่าน"],
]

# Build appendix elements
appendix_elems = []

# ── ภาคผนวก heading ──
appendix_elems.append(_add_page_break_elem())
appendix_elems.append(_make_para("ภาคผนวก", ref_idx=REF_CHAPTER, bold=True, alignment='center'))
appendix_elems.append(_make_para("", ref_idx=REF_BODY))

# ── ภาคผนวก ก ──
appendix_elems.append(_make_para("ภาคผนวก ก", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("โครงสร้างโปรเจกต์ (Project Structure)", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("", ref_idx=REF_BODY))

for line in project_structure.split('\n'):
    appendix_elems.append(_make_para(line, ref_idx=REF_BODY, font_size=14))

# ── ภาคผนวก ข ──
appendix_elems.append(_add_page_break_elem())
appendix_elems.append(_make_para("ภาคผนวก ข", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("โค้ดสำคัญของระบบ — Prisma Schema (prisma/schema.prisma)", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("", ref_idx=REF_BODY))

if schema_content:
    # RenderPrisma schema as one formatted code block for VS-Code-like look
    appendix_elems.append(_make_code_para(schema_content, font_name='Consolas', font_size=10, shading='F2F2F2'))

# ── ภาคผนวก ง ──
appendix_elems.append(_add_page_break_elem())
appendix_elems.append(_make_para("ภาคผนวก ง", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("README.md — คู่มือการติดตั้งและใช้งานระบบ", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("", ref_idx=REF_BODY))

if readme_content:
    # Render README as one formatted code block for VS-Code-like look
    appendix_elems.append(_make_code_para(readme_content, font_name='Consolas', font_size=10, shading='F2F2F2'))

# ── ภาคผนวก จ ──
appendix_elems.append(_add_page_break_elem())
appendix_elems.append(_make_para("ภาคผนวก จ", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("ผลการทดสอบการยอมรับระบบ (UAT) แบบละเอียด", ref_idx=REF_SECTION, bold=True))
appendix_elems.append(_make_para("", ref_idx=REF_BODY))

# Create UAT detailed table with newline replacements in step cells
uat_table_rows = []
for row in uat_detailed:
    # Replace literal backslash-n with real newlines in the "ขั้นตอน" column (index 2)
    new_row = list(row)
    new_row[2] = new_row[2].replace('\\n', '\n')
    uat_table_rows.append(new_row)

uat_table = _make_table_elem(
    ["ลำดับ", "รายการทดสอบ", "ขั้นตอนการทดสอบ", "ผลที่คาดหวัง", "ผลจริง", "ผลการทดสอบ"],
    uat_table_rows
)
appendix_elems.append(uat_table)

# Insert all appendix elements after the last bibliography entry
if bib_last:
    ref_elem = doc.paragraphs[bib_last]._element
    for elem in reversed(appendix_elems):
        ref_elem.addnext(elem)
    print(f"  Inserted {len(appendix_elems)} appendix elements after paragraph {bib_last}")
else:
    print("  WARNING: Could not find bibliography, appending at end")
    sp = body.find(qn('w:sectPr'))
    for elem in appendix_elems:
        if sp is not None:
            sp.addprevious(elem)
        else:
            body.append(elem)

# ══════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════
print(f"\nSaving to {OUT_PATH}...")
doc.save(OUT_PATH)
print("Done!")

# ── Verify ──
doc2 = Document(OUT_PATH)
print(f"\nVerification:")
print(f"  Total paragraphs: {len(doc2.paragraphs)}")
print(f"  Total tables: {len(doc2.tables)}")

# Check TOC
toc_found = False
fig_toc_found = False
tbl_toc_found = False
for p in doc2.paragraphs:
    if p.text.strip() == 'สารบัญ':
        toc_found = True
    if p.text.strip() == 'สารบัญรูป':
        fig_toc_found = True
    if p.text.strip() == 'สารบัญตาราง':
        tbl_toc_found = True

print(f"  สารบัญ: {'OK' if toc_found else 'MISSING'}")
print(f"  สารบัญรูป: {'OK' if fig_toc_found else 'MISSING'}")
print(f"  สารบัญตาราง: {'OK' if tbl_toc_found else 'MISSING'}")

# Check appendices
app_kor_found = False
app_kor_found2 = False
app_ngoe_found = False
app_jo_found = False
for p in doc2.paragraphs:
    t = p.text.strip()
    if t == 'ภาคผนวก ก':
        app_kor_found = True
    if t == 'ภาคผนวก ข':
        app_kor_found2 = True
    if t == 'ภาคผนวก ง':
        app_ngoe_found = True
    if t == 'ภาคผนวก จ':
        app_jo_found = True

print(f"  ภาคผนวก ก: {'OK' if app_kor_found else 'MISSING'}")
print(f"  ภาคผนวก ข: {'OK' if app_kor_found2 else 'MISSING'}")
print(f"  ภาคผนวก ง: {'OK' if app_ngoe_found else 'MISSING'}")
print(f"  ภาคผนวก จ: {'OK' if app_jo_found else 'MISSING'}")

# Check figure numbering
print("\n  Figure captions after renumbering:")
for p in doc2.paragraphs:
    t = p.text.strip()
    if t.startswith('รูปที่ 3.'):
        print(f"    {t[:70]}")
